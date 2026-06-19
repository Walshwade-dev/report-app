from docx import Document
doc = Document('/home/ace/.gemini/antigravity/scratch/report-app/backend/generated_word.docx')

print("--- SECTIONS ---")
for i, s in enumerate(doc.sections):
    print(f"Section {i}: w={s.page_width.inches:.2f}, h={s.page_height.inches:.2f}, "
          f"H={s.header.is_linked_to_previous}, F={s.footer.is_linked_to_previous}")

print("\n--- PAGE 2 (Section 1) ---")
s = doc.sections[1]
print(f"Header: {len(s.header.paragraphs)} paragraphs")
print(f"Footer: {s.footer.paragraphs[0].text if s.footer.paragraphs else 'None'}")

p_start = 0
for p in doc.paragraphs:
    if p.text == "JUJA MOBILE DAILY REPORT 1":
        p_start += 1
        if p_start == 2:
            print("FOUND SECOND TITLE")
            break

print("--- TABLE 1 (Layout Table) ---")
try:
    t = doc.tables[1]
    print(f"Layout table rows: {len(t.rows)}, cols: {len(t.columns)}")
    print(f"Left cell width: {t.cell(0,0).width}, Right cell width: {t.cell(0,1).width}")
    
    # Nested table
    inner_t = t.cell(0,0).tables[0]
    print(f"Inner table: {len(inner_t.rows)}x{len(inner_t.columns)}")
    print(f"Inner Font Size: {inner_t.cell(0,1).paragraphs[0].runs[0].font.size.pt}")
except Exception as e:
    print(f"Error inspecting layout table: {e}")
