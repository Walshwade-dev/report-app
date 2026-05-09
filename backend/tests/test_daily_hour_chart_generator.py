from pathlib import Path

import pandas as pd
from docx import Document

from app.services.daily_hour_chart_generator import add_daily_hour_chart_section
from app.services.daily_hour_processor import (
    add_daily_totals_row,
    build_daily_hour_metrics,
)


FIXTURES_DIR = Path(__file__).parent / "fixtures"


def build_daily_hour_dataframe():
    raw_df = pd.read_csv(FIXTURES_DIR / "daily_hour.csv")
    daily_df = build_daily_hour_metrics(
        raw_df,
        report_date="02/02/2026",
        wideload_count=1,
    )
    return add_daily_totals_row(daily_df)


def test_daily_hour_chart_table_has_single_formula_header_row():
    doc = Document()
    add_daily_hour_chart_section(doc, build_daily_hour_dataframe())

    container = doc.tables[0]
    table = container.rows[0].cells[0].tables[0]
    formula_cells = [cell.text.strip() for cell in table.rows[1].cells[1:]]
    first_data_cells = [cell.text.strip() for cell in table.rows[2].cells]

    assert len(container.rows) == 1
    assert len(container.rows[0].cells) == 2
    assert len(table.rows) == 27
    assert formula_cells == ["N=(D+S)", "(M)", "Q = H-C", "X= (N+M)"]
    assert first_data_cells[0] == "0000-0100"
    assert first_data_cells[1:] != formula_cells
    assert table.rows[-1].cells[0].text.strip() == "Total"
