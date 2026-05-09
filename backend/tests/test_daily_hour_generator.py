from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

from app.services.daily_hour_generator import add_daily_hour_statistics_section
from app.services.daily_hour_processor import (
    add_daily_totals_row,
    build_daily_hour_metrics,
)


FIXTURES_DIR = Path(__file__).parent / "fixtures"


def build_daily_hour_dataframe():
    raw_df = pd.read_csv(FIXTURES_DIR / "daily_hour.csv")
    daily_df = build_daily_hour_metrics(
        raw_df,
        report_date="02/02/2026",
        wideload_count=1,
    )
    return add_daily_totals_row(daily_df)


def first_run(cell):
    return cell.paragraphs[0].runs[0]


def test_daily_hour_statistics_data_rows_are_taller_with_bold_date_time_cells():
    doc = Document()
    add_daily_hour_statistics_section(doc, build_daily_hour_dataframe())

    table = doc.tables[0]
    first_data_row = table.rows[3]
    date_run = first_run(first_data_row.cells[0])
    time_run = first_run(first_data_row.cells[1])

    assert first_data_row.height == Pt(12.6)
    assert date_run.bold is False
    assert time_run.bold is True
    assert date_run.font.size == Pt(9)
    assert time_run.font.size == Pt(8)
