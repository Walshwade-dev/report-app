import io
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.enum.section import WD_ORIENT
from openpyxl import load_workbook

from app.services.report_layout import (
    A4_LANDSCAPE_HEIGHT_INCHES,
    A4_LANDSCAPE_WIDTH_INCHES,
    A4_PRINTABLE_WIDTH_TWIPS,
    BOTTOM_MARGIN_INCHES,
    LEFT_MARGIN_INCHES,
    RIGHT_MARGIN_INCHES,
    TOP_MARGIN_INCHES,
)


FIXTURES_DIR = Path(__file__).parent / "fixtures"
WORD_NAMESPACE = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
CHART_NAMESPACE = {
    "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
}
SPREADSHEET_NAMESPACE = {
    "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
}


def fixture_upload(name: str, content_type: str = "text/csv"):
    path = FIXTURES_DIR / name
    return {
        "file": (
            path.name,
            path.read_bytes(),
            content_type,
        )
    }


def create_report_session(client) -> str:
    response = client.post(
        "/api/report-sessions",
        json={
            "report_date": "2026-02-02",
            "station": "Juja",
            "bound": "Thika Bound",
            "weighbridge_name": "Juja",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["report_id"]
    return payload["report_id"]


def patch_manual_inputs(client, report_id: str) -> None:
    response = client.patch(
        f"/api/report-sessions/{report_id}/manual-inputs",
        json={
            "prepared_by": "Fredrick Kariuki",
            "confirmed_by": "Faith Njani",
            "traffic_census": {
                "buses_gte_3500kg": 1351,
                "vehicles_3500_to_7000_excluding_buses": 29,
                "vehicles_gte_7000_excluding_buses": 6,
            },
            "transgressions": {
                "daily_transgressions": [],
                "action_report": [],
            },
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["sections"]["traffic_census"]["status"] == "ready"
    assert payload["sections"]["transgressions"]["status"] == "ready"


def upload_required_files(client, report_id: str) -> None:
    wideload = client.post(
        f"/api/report-sessions/{report_id}/uploads/wideload",
        files=fixture_upload("wideload.csv"),
    )
    assert wideload.status_code == 200
    wideload_payload = wideload.json()
    assert wideload_payload["sections"]["wideload"]["status"] == "ready"
    assert wideload_payload["sections"]["wideload"]["wideload_count"] == 1

    daily_hour = client.post(
        f"/api/report-sessions/{report_id}/uploads/daily-hour",
        data={"wideload_count": "999"},
        files=fixture_upload("daily_hour.csv"),
    )
    assert daily_hour.status_code == 200
    daily_hour_payload = daily_hour.json()
    assert daily_hour_payload["sections"]["daily_hour"]["status"] == "ready"
    assert daily_hour_payload["sections"]["daily_hour"]["wideload_count_used"] == 1

    impounded = client.post(
        f"/api/report-sessions/{report_id}/uploads/impounded-prohibited",
        files=fixture_upload("impounded_prohibited.csv"),
    )
    assert impounded.status_code == 200
    assert impounded.json()["sections"]["impounded_prohibited"]["status"] == "ready"

    overloaded = client.post(
        f"/api/report-sessions/{report_id}/uploads/overloaded",
        files=fixture_upload("overloaded.csv"),
    )
    assert overloaded.status_code == 200
    overloaded_payload = overloaded.json()
    assert overloaded_payload["sections"]["overloaded"]["status"] == "ready"
    assert overloaded_payload["sections"]["overloaded"]["valid_permit_count"] == 1


def test_mobile_report_upload_extracts_weighbridge_register(client):
    report_id = create_report_session(client)

    response = client.post(
        f"/api/report-sessions/{report_id}/uploads/mobile-report",
        files=fixture_upload("mobile_report.csv"),
    )

    assert response.status_code == 200
    payload = response.json()
    summary = payload["sections"]["mobile_report"]["summary"]

    assert payload["sections"]["mobile_report"]["status"] == "ready"
    assert summary["total_records"] == 4
    assert summary["report_date"] == "2026-05-12"
    assert summary["station"] == "JJM2"
    assert summary["overloaded_records"] == 2
    assert payload["mobile_report"]["data"][0]["registration"] == "KBW781J"
    assert payload["mobile_report"]["data"][0]["total_gvw_kg"] == 37500
    assert payload["mobile_report"]["data"][0]["gvw_difference_kg"] == 7500
    assert payload["mobile_report"]["data"][0]["remarks"] == "CHARGED"
    assert payload["mobile_excel_report"]["status"] == "ready"
    assert payload["mobile_word_report"]["status"] == "ready"


def test_mobile_report_downloads_mapped_excel_workbook(client):
    report_id = create_report_session(client)

    manual_response = client.patch(
        f"/api/report-sessions/{report_id}/manual-inputs",
        json={
            "extra": {
                "mobile_report": {
                    "route": "westlands-parklands-ruiru-juja",
                    "danka_staff": "dm duncan odhiambo",
                    "police_officers": "cpl emason sautet",
                    "mobile_vehicle": "kds042z",
                    "mileage_start": 61267,
                    "mileage_end": 61447,
                    "cases_cleared_in_court": 3,
                }
            }
        },
    )
    assert manual_response.status_code == 200

    upload = client.post(
        f"/api/report-sessions/{report_id}/uploads/mobile-report",
        files=fixture_upload("mobile_report.csv"),
    )
    assert upload.status_code == 200

    download = client.get(
        f"/api/report-sessions/{report_id}/download-mobile-excel-report"
    )
    assert download.status_code == 200
    assert (
        download.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    assert download.headers["content-disposition"].endswith('.xlsx"')

    workbook = load_workbook(io.BytesIO(download.content), data_only=False)
    summary = workbook["Weigh & Hourly Summary"]
    detail = workbook["Mobile Daily Report"]

    assert summary["C5"].value.strftime("%d-%b-%y") == "12-May-26"
    assert summary["E5"].value == 1
    assert summary["G5"].value == 1
    assert summary["H5"].value == 7500
    assert summary["E11"].value == 3
    assert summary["F11"].value == 1
    assert summary["H33"].value == 3
    assert len(summary._charts) == 1
    assert summary["E5"].fill.fill_type is None
    assert summary["D32"].fill.fill_type is None

    assert detail["B6"].value.strftime("%d-%b-%y") == "12-May-26"
    assert detail["I6"].value == 4
    assert detail["J6"].value == 1
    assert detail["J7"].value == 0
    assert detail["K6"].value == 61267
    assert detail["L6"].value == 61447
    assert detail["M6"].value == "=L6-K6"
    assert detail["N6"].value == "KDS042Z"

    assert detail["C11"].value == "KBW781J"
    assert detail["D11"].value == "OMAR HALEN."
    assert detail["F11"].value == 7500
    assert detail["G11"].value == 7500
    assert detail["M11"].value == "CHARGED"
    assert detail["N11"].value == "WESTLANDS-PARKLANDS-RUIRU-JUJA"
    assert detail["B11"].number_format == "dd-mmm-yy"
    assert detail["B5"].fill.fill_type is None
    assert detail["B10"].fill.fill_type is None
    assert detail["B10"].alignment.horizontal == "center"
    assert detail["E6"].alignment.horizontal == "left"
    assert detail["E6"].alignment.wrap_text is not True
    assert detail["G6"].alignment.horizontal == "left"
    assert detail["G6"].alignment.wrap_text is not True
    assert detail["K11"].alignment.horizontal == "left"
    assert detail["K11"].alignment.wrap_text is not True
    assert detail["L11"].alignment.horizontal == "left"
    assert detail["L11"].alignment.wrap_text is not True
    assert detail["N11"].alignment.horizontal == "left"
    assert detail["N11"].alignment.wrap_text is not True
    assert detail.row_dimensions[1].height == 15.75
    assert detail.row_dimensions[3].height == 14.25
    assert detail.row_dimensions[4].height == 18.75
    assert detail.row_dimensions[7].height == 15.75
    assert detail.row_dimensions[8].height == 15.75

    assert "N3:X30" in {str(range_) for range_ in summary.merged_cells.ranges}
    assert "O32:V32" in {str(range_) for range_ in summary.merged_cells.ranges}
    assert summary["O33"].value == "Red = formulae, do not edit"
    assert summary["O33"].fill.fgColor.rgb == "FFFF0000"
    assert summary["O33"].alignment.horizontal == "left"
    chart = summary._charts[0]
    assert chart.x_axis.axPos == "b"
    assert chart.y_axis.crossBetween == "between"
    assert chart.dataLabels.showVal is False
    assert chart.dataLabels.showCatName is False
    assert chart.dataLabels.showSerName is False
    assert [series.smooth for series in chart.ser] == [False, False]
    assert [series.graphicalProperties.line.width for series in chart.ser] == [
        28575,
        28575,
    ]
    assert [series.graphicalProperties.line.cap for series in chart.ser] == [
        "rnd",
        "rnd",
    ]
    assert [series.graphicalProperties.line.solidFill.srgbClr for series in chart.ser] == [
        "1F4E79",
        "800000",
    ]

    with zipfile.ZipFile(io.BytesIO(download.content)) as archive:
        chart_xml = ElementTree.fromstring(archive.read("xl/charts/chart1.xml"))
        drawing_xml = ElementTree.fromstring(archive.read("xl/drawings/drawing1.xml"))
        summary_xml = ElementTree.fromstring(archive.read("xl/worksheets/sheet1.xml"))

    line_elements = chart_xml.findall(".//c:ser/c:spPr/a:ln", CHART_NAMESPACE)
    assert len(line_elements) == 2
    assert [line.attrib["w"] for line in line_elements] == ["28575", "28575"]
    assert [line.attrib["cap"] for line in line_elements] == ["rnd", "rnd"]
    assert all(line.find("a:round", CHART_NAMESPACE) is not None for line in line_elements)
    color_values = [
        color.attrib["val"]
        for color in chart_xml.findall(".//c:ser/c:spPr/a:ln/a:solidFill/a:srgbClr", CHART_NAMESPACE)
    ]
    assert color_values == ["1F4E79", "800000"]
    smooth_values = [
        smooth.attrib["val"]
        for smooth in chart_xml.findall(".//c:ser/c:smooth", CHART_NAMESPACE)
    ]
    assert smooth_values == ["0", "0"]
    data_labels = chart_xml.find(".//c:dLbls", CHART_NAMESPACE)
    assert data_labels is not None
    for tag in [
        "showLegendKey",
        "showVal",
        "showCatName",
        "showSerName",
        "showPercent",
        "showBubbleSize",
    ]:
        element = data_labels.find(f"c:{tag}", CHART_NAMESPACE)
        assert element is not None
        assert element.attrib["val"] == "0"
    category_axis_text = chart_xml.find(".//c:catAx/c:txPr", CHART_NAMESPACE)
    assert category_axis_text is not None
    body_pr = category_axis_text.find("a:bodyPr", CHART_NAMESPACE)
    assert body_pr is not None
    assert body_pr.attrib["rot"] == "-60000000"
    chart_anchor_to = drawing_xml.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing}graphicFrame/.."
        "/{http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing}to"
    )
    assert chart_anchor_to is not None
    chart_to_column = chart_anchor_to.find(
        "{http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing}col"
    )
    assert chart_to_column is not None
    assert chart_to_column.text == "23"

    # Verify notes cells via openpyxl (simpler than XML parsing for merged cells)
    assert summary["O33"].value is not None
    assert "Red" in str(summary["O33"].value)
    assert "O34:V34" in {str(r) for r in summary.merged_cells.ranges}
    assert "O35:V35" in {str(r) for r in summary.merged_cells.ranges}


def test_mobile_report_downloads_mapped_word_document(client):
    report_id = create_report_session(client)

    manual_response = client.patch(
        f"/api/report-sessions/{report_id}/manual-inputs",
        json={
            "prepared_by": "Fredrick Kariuki",
            "confirmed_by": "Faith Njani",
            "extra": {
                "mobile_report": {
                    "route": "westlands-parklands-ruiru-juja",
                    "danka_staff": "dm duncan odhiambo",
                    "police_officers": "cpl emason sautet",
                    "mobile_vehicle": "kds042z",
                    "mileage_start": 61267,
                    "mileage_end": 61447,
                    "cases_cleared_in_court": 3,
                }
            },
        },
    )
    assert manual_response.status_code == 200

    upload = client.post(
        f"/api/report-sessions/{report_id}/uploads/mobile-report",
        files=fixture_upload("mobile_report.csv"),
    )
    assert upload.status_code == 200
    assert upload.json()["mobile_word_report"]["download_url"].endswith(
        "/download-mobile-word-report"
    )

    download = client.get(
        f"/api/report-sessions/{report_id}/download-mobile-word-report"
    )

    assert download.status_code == 200
    assert (
        download.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert download.headers["content-disposition"].endswith('.docx"')
    assert download.content.startswith(b"PK")

    document = Document(io.BytesIO(download.content))

    assert len(document.sections) == 10
    assert document.sections[0].orientation == WD_ORIENT.PORTRAIT
    assert all(
        section.orientation == WD_ORIENT.LANDSCAPE
        for section in document.sections[1:]
    )
    assert len(document.tables) == 9

    assert any(
        paragraph.text == "DAILY AND HOURLY STATISTICS"
        for paragraph in document.paragraphs
    )
    assert any(
        paragraph.text == "Prepared by: Fredrick Kariuki"
        for paragraph in document.paragraphs
    )
    assert any(
        paragraph.text == "Approved by: Faith Njani"
        for paragraph in document.paragraphs
    )

    daily_stats = document.tables[0]
    assert len(daily_stats.rows) == 27
    assert len(daily_stats.columns) == 6
    assert daily_stats.cell(0, 0).text == "Date"
    assert daily_stats.cell(2, 0).text == "12-MAY-26"
    assert daily_stats.cell(2, 1).text == "0000-0100"
    assert daily_stats.cell(2, 2).text == "1"
    assert daily_stats.cell(2, 4).text == "1"
    assert daily_stats.cell(26, 1).text == "Total"

    hourly_data = document.tables[1]
    assert len(hourly_data.rows) == 26
    assert len(hourly_data.columns) == 3
    assert hourly_data.cell(0, 1).text == "WEIGHED"
    assert hourly_data.cell(1, 0).text == "0000-0100"

    daily_summary = document.tables[2]
    assert daily_summary.cell(0, 0).text == "Total Weighed (X)"
    assert daily_summary.cell(1, 0).text == "4"
    assert daily_summary.cell(1, 4).text == "3"

    details = document.tables[6]
    assert len(details.rows) == 6
    assert details.cell(0, 0).text == "DATE WEIGHED"
    assert details.cell(0, 4).text == "EXCESS   WEIGHT"
    assert details.cell(1, 4).text == "GVW"
    assert details.cell(2, 0).text == "12-May-26"
    assert details.cell(2, 1).text == "KBW781J"
    assert details.cell(2, 2).text == "OMAR HALEN."
    assert details.cell(2, 4).text == "7,500"
    assert details.cell(2, 9).text == "DM DUNCAN ODHIAMBO"
    assert details.cell(2, 10).text == "CPL EMASON SAUTET"

    charged_over_two = document.tables[7]
    assert len(charged_over_two.rows) == 3
    assert charged_over_two.cell(2, 1).text == "KBW781J"

    mileage = document.tables[8]
    assert mileage.cell(0, 0).text == "MILEAGE START"
    assert mileage.cell(1, 0).text == "61,267"
    assert mileage.cell(1, 1).text == "61,447"
    assert mileage.cell(1, 2).text == "180"
    assert mileage.cell(1, 3).text == "KDS042Z"

    assert any(
        "ACTUAL ROUTE:" in paragraph.text
        and "WESTLANDS-PARKLANDS-RUIRU-JUJA" in paragraph.text
        for paragraph in document.paragraphs
    )

    with zipfile.ZipFile(io.BytesIO(download.content)) as archive:
        media_files = [
            name for name in archive.namelist() if name.startswith("word/media/")
        ]
    assert media_files


def assert_docx_preview(client, report_id: str, section_name: str) -> None:
    response = client.get(
        f"/api/report-sessions/{report_id}/sections/{section_name}/preview?format=docx"
    )

    assert response.status_code == 200
    assert response.content.startswith(b"PK")
    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def docx_traffic_census_values(docx_bytes: bytes) -> dict[str, str]:
    document = Document(io.BytesIO(docx_bytes))

    for table in document.tables:
        if len(table.rows) < 2 or len(table.rows[0].cells) < 7:
            continue

        headers = [
            cell.text.replace("\n", " ").strip()
            for cell in table.rows[0].cells
        ]

        if not any("Exemption" in header and "(E)" in header for header in headers):
            continue

        values = [cell.text.strip() for cell in table.rows[1].cells]
        return {
            "k": values[3],
            "e": values[4],
            "x": values[5],
            "total_traffic": values[6],
        }

    raise AssertionError("Traffic Census table was not found in generated DOCX.")


def expected_traffic_census_values(temp_store, report_id: str) -> dict[str, str]:
    session = temp_store.require(report_id)
    daily_df = session.dataframes["daily_hour"]
    totals_row = daily_df[
        daily_df["DATE"].astype(str).str.strip().str.lower().eq("totals")
    ].iloc[-1]

    q = int(totals_row["Q"])
    x = int(totals_row["X"])
    e = int(session.sections["wideload"]["wideload_count"])
    k = int(session.manual_inputs["traffic_census"]["total_traffic_census"])

    return {
        "k": f"{k:,}",
        "e": f"{e:,}",
        "x": f"{x:,}",
        "total_traffic": f"{q + x + k + e:,}",
    }


def expected_summary_card_values(temp_store, report_id: str) -> dict[str, int]:
    session = temp_store.require(report_id)
    daily_df = session.dataframes["daily_hour"]
    totals_row = daily_df[
        daily_df["DATE"].astype(str).str.strip().str.lower().eq("totals")
    ].iloc[-1]

    return {
        "Total Weighed": int(totals_row["X"]),
        "Total Overloaded": int(totals_row["Y"]),
        "Special Released": int(totals_row["G"]),
        "Wide Loads": int(session.sections["wideload"]["wideload_count"]),
    }


def assert_a4_landscape_layout(docx_bytes: bytes) -> None:
    document = Document(io.BytesIO(docx_bytes))

    for section in document.sections:
        assert round(section.page_width.inches, 2) == A4_LANDSCAPE_WIDTH_INCHES
        assert round(section.page_height.inches, 2) == A4_LANDSCAPE_HEIGHT_INCHES
        assert round(section.left_margin.inches, 2) == LEFT_MARGIN_INCHES
        assert round(section.right_margin.inches, 2) == RIGHT_MARGIN_INCHES
        assert round(section.top_margin.inches, 2) == TOP_MARGIN_INCHES
        assert round(section.bottom_margin.inches, 2) == BOTTOM_MARGIN_INCHES


def assert_table_grids_fit_a4_printable_width(docx_bytes: bytes) -> None:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
        document_xml = docx_zip.read("word/document.xml")

    root = ElementTree.fromstring(document_xml)

    for grid in root.findall(".//w:tblGrid", WORD_NAMESPACE):
        widths = [
            int(column.attrib[f"{{{WORD_NAMESPACE['w']}}}w"])
            for column in grid.findall("w:gridCol", WORD_NAMESPACE)
        ]
        assert sum(widths) <= A4_PRINTABLE_WIDTH_TWIPS


def get_asset_path(filename):
    p = Path("backend/app/assets") / filename
    if not p.exists():
        p = Path("app/assets") / filename
    return p


def assert_docx_uses_bgwhite_logo(docx_bytes: bytes) -> None:
    expected_logo = get_asset_path("bgwhitelogo.png").read_bytes()
    app_logo = get_asset_path("logo.png").read_bytes()

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
        media_files = [
            docx_zip.read(name)
            for name in docx_zip.namelist()
            if name.startswith("word/media/")
        ]

    assert expected_logo in media_files
    assert app_logo not in media_files


def test_upload_fixtures_build_and_download_final_report(client, temp_store):
    report_id = create_report_session(client)
    upload_required_files(client, report_id)
    patch_manual_inputs(client, report_id)

    session_payload = client.get(f"/api/report-sessions/{report_id}").json()
    assert session_payload["sections"]["daily_summary"]["status"] == "ready"
    expected_traffic_values = expected_traffic_census_values(temp_store, report_id)
    expected_summary_values = expected_summary_card_values(temp_store, report_id)

    summary_cards = client.get(f"/api/report-sessions/{report_id}/summary-cards")
    assert summary_cards.status_code == 200
    summary_payload = summary_cards.json()
    assert summary_payload["report_id"] == report_id
    assert {
        card["title"]: card["value"]
        for card in summary_payload["cards"]
    } == expected_summary_values
    assert all(card["status"] == "ready" for card in summary_payload["cards"])

    for section_name in [
        "daily-hour-statistics",
        "traffic-census",
        "daily-summary",
        "transgressions",
        "wideload",
        "impounded-prohibited",
    ]:
        assert_docx_preview(client, report_id, section_name)

    traffic_preview = client.get(
        f"/api/report-sessions/{report_id}/sections/traffic-census/preview?format=docx"
    )
    assert traffic_preview.status_code == 200
    assert docx_traffic_census_values(traffic_preview.content) == expected_traffic_values
    assert_docx_uses_bgwhite_logo(traffic_preview.content)

    build_response = client.post(f"/api/report-sessions/{report_id}/build-final-report")
    assert build_response.status_code == 200
    build_payload = build_response.json()
    assert build_payload["final_report"]["status"] == "ready"

    final_report_path = temp_store.final_reports_dir / report_id / "final_report.docx"
    assert final_report_path.exists()

    download = client.get(f"/api/report-sessions/{report_id}/download-final-report")
    assert download.status_code == 200
    assert download.content.startswith(b"PK")
    assert download.content == final_report_path.read_bytes()
    assert docx_traffic_census_values(download.content) == expected_traffic_values
    assert_a4_landscape_layout(download.content)
    assert_table_grids_fit_a4_printable_width(download.content)
    assert_docx_uses_bgwhite_logo(download.content)
    assert (
        download.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_upload_fixtures_download_excel_report(client, temp_store):
    report_id = create_report_session(client)
    upload_required_files(client, report_id)
    patch_manual_inputs(client, report_id)

    session_payload = client.get(f"/api/report-sessions/{report_id}").json()
    assert session_payload["excel_report"]["download_url"] == (
        f"/api/report-sessions/{report_id}/download-excel-report"
    )

    download = client.get(f"/api/report-sessions/{report_id}/download-excel-report")

    assert download.status_code == 200
    assert (
        download.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    workbook = load_workbook(io.BytesIO(download.content))
    assert workbook.sheetnames == ["Summary", "CC records", "Hswim Weekly"]

    summary = workbook["Summary"]
    assert summary["D2"].value == "Trucks Weighed"
    assert summary["T2"].value == "TIME"
    assert summary["C33"].value == "Buses>= 3500kg  (J)"
    assert summary["H33"].value == "Exemption\npermits"
    assert summary["H34"].value == "(E)"
    assert summary["I33"].value == "Total Weighed"
    assert summary["I34"].value == "(X)"
    assert summary["J33"].value == "Total Traffic"
    assert summary["J34"].value == "(T)= (Q+X+K+E)"
    assert summary["B38"].value == "HSWIM CLEARED (Q)"
    assert len(summary._charts) == 1
    assert summary.column_dimensions["B"].width == 12.140625
    assert summary.row_dimensions[38].height == 67.5
    assert summary["D6"].fill.fgColor.rgb == "FFFFFF00"
    assert summary["H6"].fill.fgColor.rgb == "FFFF0000"
    assert summary["N40"].fill.fgColor.rgb == "00000000"
    assert summary["O40"].fill.fgColor.rgb == "00000000"
    assert summary["Q40"].fill.fgColor.rgb == "00000000"
    assert summary["P40"].value == "=Q30"
    assert summary._charts[0].x_axis.axPos == "b"
    assert "H34:H35" in {str(merged) for merged in summary.merged_cells.ranges}
    assert "I34:I35" in {str(merged) for merged in summary.merged_cells.ranges}
    assert "J34:J35" in {str(merged) for merged in summary.merged_cells.ranges}
    assert summary._charts[0].y_axis.scaling.min == 0
    assert summary._charts[0].y_axis.scaling.max == 300
    assert summary._charts[0].y_axis.majorUnit == 50
    assert summary._charts[0].x_axis.crosses == "min"
    assert all(series.smooth is False for series in summary._charts[0].series)
    assert [
        series.graphicalProperties.line.solidFill.srgbClr
        for series in summary._charts[0].series
    ] == ["4472C4", "ED7D31", "A5A5A5", "FFC000"]
    assert [
        series.graphicalProperties.line.width
        for series in summary._charts[0].series
    ] == [38100, 38100, 38100, 38100]
    assert summary["S48"].value == "Notes"
    assert summary["S48"].font.bold is True
    assert summary["S48"].font.sz == 20
    assert summary["S48"].alignment.horizontal == "center"
    assert summary["S48"].border.top.style == "medium"
    assert summary["S49"].fill.fgColor.rgb == "FFFFFF00"
    assert summary["S50"].fill.fgColor.rgb == "FFFF0000"
    assert summary["S51"].fill.fill_type is None
    assert summary["S52"].value == "Data in the table is for illustration purposes ONLY"
    assert summary["S52"].border.top.style is None

    cc_records = workbook["CC records"]
    assert cc_records["B2"].value == "Total Traffic Census Summary"
    assert "B8:D10" not in {str(merged) for merged in cc_records.merged_cells.ranges}
    assert cc_records["B8"].value is not None
    assert cc_records["C9"].value is not None
    assert cc_records["D10"].value is not None
    assert cc_records["B14"].value == "=SUM(B6:B13)"
    assert cc_records["C14"].value == "=SUM(C6:C13)"
    assert cc_records["D14"].value == "=SUM(D6:D13)"
    assert cc_records["B14"].fill.fgColor.rgb == "FF00B050"
    assert cc_records["C14"].fill.fgColor.rgb == "FF00B050"
    assert cc_records["D14"].fill.fgColor.rgb == "FF00B050"


def test_traffic_census_preview_requires_wideload_upload(client, temp_store):
    report_id = create_report_session(client)
    patch_manual_inputs(client, report_id)

    response = client.get(
        f"/api/report-sessions/{report_id}/sections/traffic-census/preview?format=docx"
    )

    assert response.status_code == 400
    assert response.json()["detail"] == (
        "Traffic Census preview requires wideload upload first because E comes from wideload count."
    )


def test_summary_cards_report_awaiting_data_for_new_session(client, temp_store):
    report_id = create_report_session(client)

    response = client.get(f"/api/report-sessions/{report_id}/summary-cards")

    assert response.status_code == 200
    payload = response.json()
    assert [card["title"] for card in payload["cards"]] == [
        "Total Weighed",
        "Total Overloaded",
        "Special Released",
        "Wide Loads",
    ]
    assert all(card["value"] is None for card in payload["cards"])
    assert all(card["display_value"] == "—" for card in payload["cards"])
    assert all(card["status"] == "awaiting_data" for card in payload["cards"])
