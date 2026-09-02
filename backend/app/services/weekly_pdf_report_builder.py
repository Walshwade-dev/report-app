import io
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.services.preview_renderer import convert_docx_to_pdf

def set_cell_borders(cell, outer_only=False, is_first_col=False, is_last_col=False):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = ['top', 'bottom']
    if not outer_only:
        borders.extend(['left', 'right'])
    else:
        if is_first_col:
            borders.append('left')
        if is_last_col:
            borders.append('right')
            
    for border_name in borders:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    
    tcPr.append(tcBorders)

def set_cell_background(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_table_width(table):
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    
    # Force fixed layout to ensure table doesn't overflow margins
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

def _build_table_in_doc(doc, title_text: str, columns: list[str], data: list[dict]):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(12)

    table = doc.add_table(rows=2, cols=len(columns))
    table.style = 'Table Grid'
    set_table_width(table)

    hdr_cells = table.rows[0].cells
    for i in range(17):
        hdr_cells[i].text = columns[i]
        hdr_cells[i].merge(table.rows[1].cells[i])
    
    hdr_cells[17].text = "Exemption Permits"
    hdr_cells[17].merge(hdr_cells[19])
    
    sub_headers = ["Not Weighd (E)", "Weighed(F)", "Total"]
    for i in range(3):
        table.rows[1].cells[17 + i].text = sub_headers[i]

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in paragraph.runs:
                    r.bold = True
                    r.font.size = Pt(7)

    for row_data in data:
        row_cells = table.add_row().cells
        for idx, col in enumerate(columns):
            val = row_data.get(col, "")
            row_cells[idx].text = str(val) if pd.notna(val) else ""
            row_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in row_cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in paragraph.runs:
                    r.font.size = Pt(7)
    
    df = pd.DataFrame(data, columns=columns)
    totals_row = table.add_row().cells
    totals_row[0].text = "Total"
    for idx, col in enumerate(columns[1:], start=1):
        total_val = pd.to_numeric(df[col], errors='coerce').fillna(0).sum()
        total_val = int(total_val) if pd.notnull(total_val) else ""
        totals_row[idx].text = str(total_val)
        totals_row[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in totals_row[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in paragraph.runs:
                r.bold = True
                r.font.size = Pt(7)

    for paragraph in totals_row[0].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in paragraph.runs:
            r.bold = True
            r.font.size = Pt(7)

    doc.add_paragraph()

def build_weekly_pdf_report(
    weekly_data_by_bound: dict[str, list[dict]],
    weekly_data_combined: list[dict],
    start_date: str,
    end_date: str,
    station: str,
    prepared_by: str,
    approved_by: str,
) -> io.BytesIO:
    doc = Document()
    
    # Set landscape Letter (11 x 8.5)
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    
    # Add Logo
    logo_path = os.path.join(os.path.dirname(__file__), "../assets/logo.png")
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_para.add_run().add_picture(logo_path, width=Inches(2))

    columns = [
        "DATE", "HSWIM Total (H)", "Called in (C)", "Cleared by HSWIM\n(Q) = (H-C)",
        "Weighed Scale (N)= D+S", "Manually Weighed (M)", "Total Weighed (X) = (N+M)",
        "Total Traffic Census =(K)", "Total Traffic (T) = (Q+X+K+E)", "Total Overloaded (Y)=(A+G+P)",
        "Impounded & Prohibited (P) = (Z+R)", "Warned (A)", "Prohibited & Charged (Z)",
        "Special Released (G)", "Redistributed (R)", "Cases Cleared in Court (B)",
        "Transgressions (L)", "Not Weighd (E)", "Weighed(F)", "Total"
    ]

    for bound_name, data in weekly_data_by_bound.items():
        title = f"{station.upper()} WEIGHBRIDGE {bound_name.upper()} WEEKLY SUMMARY REPORT"
        _build_table_in_doc(doc, title, columns, data)

    title_combined = f"{station.upper()} WEIGHBRIDGE WEEKLY TOTAL SUMMARY REPORT"
    _build_table_in_doc(doc, title_combined, columns, weekly_data_combined)

    # Signatures
    sigs_para = doc.add_paragraph()
    sigs_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_prep = sigs_para.add_run(f"PREPARED BY: {prepared_by.upper()}")
    run_prep.bold = True
    run_prep.font.name = "Calibri"
    run_prep.font.size = Pt(14)
    
    doc.add_paragraph()
    
    sigs_para2 = doc.add_paragraph()
    sigs_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_app = sigs_para2.add_run(f"APPROVED BY: {approved_by.upper()}")
    run_app.bold = True
    run_app.font.name = "Calibri"
    run_app.font.size = Pt(14)

    doc.add_paragraph()

    # Footer Table
    footer_table = doc.add_table(rows=1, cols=3)
    set_table_width(footer_table)
    
    f_cells = footer_table.rows[0].cells
    f_cells[0].text = f"{station.upper()} WEIGHBRIDGE WEEKLY SUMMARY REPORT"
    f_cells[1].text = "KeNHA/WB/MTCE/4339/2025"
    f_cells[2].text = f"FOR DATES: {start_date} TO {end_date}"

    # Apply borders and light grey background
    for i, cell in enumerate(f_cells):
        is_first = (i == 0)
        is_last = (i == len(f_cells) - 1)
        set_cell_borders(cell, outer_only=True, is_first_col=is_first, is_last_col=is_last)
        set_cell_background(cell, "EEEEEE")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(14)

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    pdf_content, _ = convert_docx_to_pdf(docx_buffer, "weekly_report.docx")
    
    return io.BytesIO(pdf_content)
