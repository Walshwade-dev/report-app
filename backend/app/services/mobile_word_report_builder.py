import io
import re
from datetime import datetime
from typing import Any
from xml.sax.saxutils import escape

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.services.daily_hour_processor import HOURS
from app.services.mobile_report_processor import summarize_mobile_report
from app.services.report_layout import LOGO_PATH, enable_field_updates


FONT_NAME = "Arial"
DARK_BLUE_LINE = "#1F4E79"
MAROON_LINE = "#800000"
LETTER_WIDTH_INCHES = 8.5
LETTER_HEIGHT_INCHES = 11.0
LANDSCAPE_WIDTH_INCHES = 11.0
LANDSCAPE_HEIGHT_INCHES = 8.5
VEHICLE_TABLE_WIDTHS = [
    1313,
    1155,
    1300,
    656,
    774,
    774,
    1313,
    1313,
    1076,
    1969,
    1562,
    1195,
]


def _manual_source(session) -> dict[str, Any]:
    mobile = session.manual_inputs.get("mobile_report")
    if isinstance(mobile, dict):
        return mobile
    return {}


def _manual_value(session, *keys: str, default: Any = "") -> Any:
    mobile = _manual_source(session)
    for source in (mobile, session.manual_inputs):
        for key in keys:
            value = source.get(key)
            if value not in (None, ""):
                return value
    return default


def _manual_int(session, *keys: str, default: int = 0) -> int:
    value = _manual_value(session, *keys, default=default)
    number = pd.to_numeric(value, errors="coerce")
    if pd.isna(number):
        return default
    return int(number)


def _upper(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return text.upper() if text else ""


def _mobile_station_label(session) -> str:
    station = _upper(getattr(session, "station", "JUJA"))
    station = re.sub(r"\bMOBILE\b", "", station).strip()
    station = re.sub(r"\s+", " ", station)
    if not station or station == "NONE":
        return "JUJA"
    return station


def _mobile_report_number(session) -> str:
    bound = str(getattr(session, "bound", "") or "").strip().lower()
    if re.search(r"\b(2|two|mobile_2)\b", bound):
        return "2"
    return "1"


def _date_value(value: Any) -> datetime:
    if isinstance(value, datetime):
        return value
    return pd.to_datetime(value).to_pydatetime()


def _display_date(value: Any, *, upper: bool = False) -> str:
    date_value = _date_value(value)
    month_name = date_value.strftime("%B")
    month_label = month_name if len(month_name) <= 4 else month_name[:3]
    if upper:
        month_label = month_label.upper()
    return f"{date_value.day:02d}-{month_label}-{date_value:%y}"


def _vehicle_display_date(value: Any) -> str:
    date_value = _date_value(value)
    return f"{date_value.day:02d}-{date_value.strftime('%B')[:3].upper()}-{date_value:%y}"


def _format_number(value: Any) -> str:
    number = pd.to_numeric(value, errors="coerce")
    if pd.isna(number):
        return "-"
    number = int(number)
    if number <= 0:
        return "-"
    return f"{number:,}"


def _format_plain_number(value: Any) -> str:
    number = pd.to_numeric(value, errors="coerce")
    if pd.isna(number):
        return ""
    return f"{int(number)}"


def _plain_int(value: Any) -> int | None:
    number = pd.to_numeric(str(value).replace(",", ""), errors="coerce")
    if pd.isna(number):
        return None
    return int(number)


def _set_section_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(LETTER_WIDTH_INCHES)
    section.page_height = Inches(LETTER_HEIGHT_INCHES)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.63)
    section.right_margin = Inches(1.25)


def _set_section_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(LANDSCAPE_WIDTH_INCHES)
    section.page_height = Inches(LANDSCAPE_HEIGHT_INCHES)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.0)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def _add_landscape_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_landscape(section)
    return section


def _add_continuous_landscape_section(doc: Document):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_section_landscape(section)
    return section


def _add_initial_landscape_section(doc: Document, report_date, session):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_landscape(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _add_mobile_title_header(section, session)
    _add_mobile_footer(section, report_date, session)
    return section


def _set_table_width(table, width_twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def _set_fixed_table_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _set_table_grid(table, widths: list[int]) -> None:
    existing_grid = table._tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        table._tbl.remove(existing_grid)

    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    table._tbl.insert(0, grid)


def _set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def _set_table_borders(table, size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def _set_row_height(row, inches: float) -> None:
    row.height = Inches(inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _style_cell(
    cell,
    *,
    font_name: str = FONT_NAME,
    size: float = 8,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
) -> None:
    cell.vertical_alignment = valign
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(size)
            run.bold = bold
            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)
            r_fonts.set(qn("w:ascii"), font_name)
            r_fonts.set(qn("w:hAnsi"), font_name)


def _set_cell_text(
    cell,
    text: Any,
    *,
    font_name: str = FONT_NAME,
    size: float = 8,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
) -> None:
    cell.text = "" if text is None else str(text)
    _style_cell(cell, font_name=font_name, size=size, bold=bold, align=align, valign=valign)


def _add_heading(doc: Document, text: str, *, size: float = 12, underline: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(str(text).upper())
    run.bold = True
    if underline:
        run.underline = True
    run.font.name = "Arial"
    run.font.size = Pt(size)


def _add_bold_line(doc: Document, text: str, *, size: float = 11) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(size)


def _add_mixed_line(
    doc: Document,
    bold_prefix: str,
    normal_text: str,
    *,
    space_before: float = 0,
    space_after: float = 0,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run1 = paragraph.add_run(bold_prefix)
    run1.bold = True
    run1.font.name = FONT_NAME
    run1.font.size = Pt(11)
    
    run2 = paragraph.add_run(normal_text)
    run2.bold = False
    run2.font.name = FONT_NAME
    run2.font.size = Pt(11)


def _hour_records(records: pd.DataFrame, hour: str) -> pd.DataFrame:
    return records.loc[records["hour_band"].eq(hour)]


def _hour_values(records: pd.DataFrame, hour: str) -> dict[str, int]:
    rows = _hour_records(records, hour)
    warned = rows["remarks"].str.strip().str.upper().eq("WARNED")
    charged = rows["is_gvw_axle_charge"] | rows["is_dimension_charge"]
    charged_mask = rows["remarks"].str.strip().str.upper().eq("CHARGED")
    excess_gvw = rows.loc[charged_mask, "gvw_difference_kg"].clip(lower=0).sum()
    return {
        "weighed": int(rows["is_weighed"].sum()),
        "warned": int(warned.sum()),
        "charged": int(charged.sum()),
        "excess_gvw": int(excess_gvw),
    }


def _summary_values(records: pd.DataFrame, session) -> dict[str, int]:
    summary = summarize_mobile_report(records)
    charged = int(summary["charged_gvw_axle_trucks"]) + int(
        summary["charged_dimensions_trucks"]
    )
    warned = int(summary["warned_trucks"])
    return {
        "total_weighed": int(summary["total_trucks_weighed"]),
        "warned_and_charged": warned + charged,
        "charged": charged,
        "warned": warned,
        "cases_cleared": _manual_int(
            session,
            "cases_cleared_in_court",
            "cases_cleared_court",
            default=0,
        ),
        "transgressions": _manual_int(session, "transgressions_count", default=0),
        "exempted_permit": _manual_int(session, "exempted_permit", default=0),
        "manually_weighed": _manual_int(session, "manually_weighed", default=0),
    }


def _mobile_hourly_chart_scale(max_value: int) -> tuple[int, int]:
    if max_value <= 15:
        upper = max(2, max_value + 2)
        return upper, 2

    if max_value <= 20:
        upper = max(2, ((max_value + 1) // 2) * 2)
        return upper, 2

    return max_value + 5, 5


def _style_mobile_hourly_chart_axes(ax) -> None:
    ax.grid(axis="y", color="#D9D9D9", linewidth=0.8)
    ax.grid(axis="x", visible=False)
    ax.tick_params(
        axis="x",
        labelrotation=48,
        labelsize=6,
        colors="#595959",
        length=0,
    )
    ax.tick_params(axis="y", labelsize=8, colors="#595959", length=0)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.spines["bottom"].set_color("#BFBFBF")


def _plot_mobile_hourly_chart_lines(ax, weighed: list[int], charged: list[int]) -> None:
    ax.plot(HOURS, weighed, label="WEIGHED", color=DARK_BLUE_LINE, linewidth=2.2)
    ax.plot(HOURS, charged, label="CHARGED", color=MAROON_LINE, linewidth=2.2)


def _create_mobile_hourly_chart(records: pd.DataFrame) -> io.BytesIO:
    weighed = [_hour_values(records, hour)["weighed"] for hour in HOURS]
    charged = [_hour_values(records, hour)["charged"] for hour in HOURS]
    max_value = max([*weighed, *charged, 1])
    upper, tick_step = _mobile_hourly_chart_scale(max_value)

    buffer = io.BytesIO()
    fig = plt.figure(figsize=(6.3, 4.1), dpi=150)
    ax = fig.add_axes([0.08, 0.2, 0.9, 0.66])
    fig.patch.set_facecolor("white")
    fig.patch.set_edgecolor("#D9D9D9")
    fig.patch.set_linewidth(1.0)

    _plot_mobile_hourly_chart_lines(ax, weighed, charged)
    ax.set_title("DAILY HOURLY DATA", fontsize=13, fontweight="bold", color="#595959")
    ax.set_ylim(0, upper)
    ax.set_yticks(range(0, upper + 1, tick_step))
    _style_mobile_hourly_chart_axes(ax)
    ax.legend(
        fontsize=7,
        loc="center",
        bbox_to_anchor=(0.5, 0.06),
        bbox_transform=fig.transFigure,
        ncol=2,
        frameon=False,
    )

    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _add_daily_hour_statistics(doc: Document, records: pd.DataFrame, report_date) -> None:
    _add_heading(doc, "1.0 DAILY AND HOURLY STATISTICS", size=12, underline=False)
    table = doc.add_table(rows=27, cols=6)
    table.autofit = False
    table.allow_autofit = False
    widths = [1913, 1867, 1644, 1484, 1609, 2283]
    _set_fixed_table_layout(table)
    _set_table_grid(table, widths)
    _set_table_width(table, sum(widths))
    _set_table_borders(table)

    _set_row_height(table.rows[0], 0.3458)
    _set_row_height(table.rows[1], 0.2868)
    _set_row_height(table.rows[2], 0.3368)
    for row in table.rows[3:]:
        _set_row_height(row, 0.2736)

    headers_1 = [
        "Date",
        "Time",
        "Trucks Weighed",
        "Warned Trucks",
        "Charged & Prohibited Trucks",
        "Excess Weight",
    ]
    for col, text in enumerate(headers_1):
        _set_cell_text(table.cell(0, col), text, size=11, bold=True)
        _set_cell_width(table.cell(0, col), widths[col])
        
    _set_cell_text(table.cell(1, 5), "GVW (KGS)", size=11, bold=True)
    _set_cell_width(table.cell(1, 5), widths[5])

    for col in range(5):
        cell_0 = table.cell(0, col)
        cell_1 = table.cell(1, col)
        cell_0.merge(cell_1)
        _set_cell_text(cell_0, headers_1[col], size=11, bold=True)

    totals = {"weighed": 0, "warned": 0, "charged": 0, "excess_gvw": 0}
    for index, hour in enumerate(HOURS, start=2):
        values = _hour_values(records, hour)
        for key in totals:
            totals[key] += values[key]

        row_values = [
            _display_date(report_date, upper=True) if index == 2 else "",
            hour,
            values["weighed"],
            values["warned"],
            values["charged"],
            _format_plain_number(values["excess_gvw"]),
        ]
        for col, value in enumerate(row_values):
            _set_cell_text(table.cell(index, col), value, font_name="Calibri", size=11)
            _set_cell_width(table.cell(index, col), widths[col])

    total_row = table.rows[26]
    for col, value in enumerate([
        "",
        "Total",
        totals["weighed"],
        totals["warned"],
        totals["charged"],
        _format_plain_number(totals["excess_gvw"]),
    ]):
        _set_cell_text(total_row.cells[col], value, font_name="Calibri", size=11, bold=True)
        _set_cell_width(total_row.cells[col], widths[col])


def _add_prepared_approved(doc: Document, session) -> None:
    prepared_by = str(
        session.prepared_by
        or _manual_value(session, "prepared_by", default="")
    ).strip()
    approved_by = str(
        session.confirmed_by
        or _manual_value(session, "approved_by", "confirmed_by", default="")
    ).strip()

    if not prepared_by and not approved_by:
        return

    paragraph = parse_xml(
        f"""
        <w:p
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office">
            <w:pPr>
                <w:spacing w:before="0" w:after="120"/>
                <w:jc w:val="left"/>
            </w:pPr>
            <w:r>
                <w:pict>
                    <v:shape
                        id="prepared-approved-textbox"
                        type="#_x0000_t202"
                        style="width:260pt;height:42pt"
                        stroked="t"
                        strokecolor="#000000"
                        strokeweight="1pt"
                        filled="f"
                        o:insetmode="auto">
                        <v:textbox inset="6pt,4pt,6pt,4pt">
                            <w:txbxContent>
                                <w:p>
                                    <w:pPr>
                                        <w:jc w:val="left"/>
                                    </w:pPr>
                                    <w:r>
                                        <w:rPr>
                                            <w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>
                                            <w:b/>
                                            <w:sz w:val="22"/>
                                        </w:rPr>
                                        <w:t>Prepared by: {escape(prepared_by)}</w:t>
                                    </w:r>
                                </w:p>
                                <w:p>
                                    <w:pPr>
                                        <w:jc w:val="left"/>
                                    </w:pPr>
                                    <w:r>
                                        <w:rPr>
                                            <w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>
                                            <w:b/>
                                            <w:sz w:val="22"/>
                                        </w:rPr>
                                        <w:t>Approved by: {escape(approved_by)}</w:t>
                                    </w:r>
                                </w:p>
                            </w:txbxContent>
                        </v:textbox>
                    </v:shape>
                </w:pict>
            </w:r>
        </w:p>
        """
    )
    body = doc._body._element
    insert_index = len(body)
    if len(body) and body[-1].tag == qn("w:sectPr"):
        insert_index -= 1
    body.insert(insert_index, paragraph)


def _add_daily_hourly_data(doc: Document, records: pd.DataFrame, session) -> None:
    _add_heading(doc, "2.0 DAILY HOURLY DATA", size=12)
    
    # Borderless layout table
    layout_table = doc.add_table(rows=1, cols=2)
    layout_table.autofit = False
    _set_fixed_table_layout(layout_table)
    _set_table_grid(layout_table, [5501, 9619])
    _set_table_width(layout_table, 15120)
    
    left_cell = layout_table.cell(0, 0)
    _set_cell_width(left_cell, 5501)
    right_cell = layout_table.cell(0, 1)
    _set_cell_width(right_cell, 9619)
    
    # Left Cell: Table
    table = left_cell.add_table(rows=26, cols=3)
    widths = [1615, 1786, 2100]
    _set_fixed_table_layout(table)
    _set_table_grid(table, widths)
    _set_table_width(table, sum(widths))
    _set_table_borders(table)

    _set_row_height(table.rows[0], 0.2083)
    _set_row_height(table.rows[1], 0.2764)
    for row in table.rows[2:]:
        _set_row_height(row, 0.2083)

    for col, text in enumerate(["", "WEIGHED", "CHARGED & PROHIBITED"]):
        _set_cell_text(table.cell(0, col), text, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_width(table.cell(0, col), widths[col])

    totals = {"weighed": 0, "charged": 0}
    for index, hour in enumerate(HOURS, start=1):
        values = _hour_values(records, hour)
        totals["weighed"] += values["weighed"]
        totals["charged"] += values["charged"]
        for col, value in enumerate([hour, values["weighed"], values["charged"]]):
            _set_cell_text(table.cell(index, col), value, font_name="Calibri", size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_width(table.cell(index, col), widths[col])

    for col, value in enumerate(["Total", totals["weighed"], totals["charged"]]):
        _set_cell_text(table.cell(25, col), value, font_name="Calibri", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_width(table.cell(25, col), widths[col])

    # Right Cell: Chart
    chart_image = _create_mobile_hourly_chart(records)
    paragraph = right_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(chart_image, width=Inches(7.5))


def _add_daily_summary(doc: Document, records: pd.DataFrame, session) -> None:
    _add_heading(doc, "3.0 daily summary")
    table = doc.add_table(rows=2, cols=8)
    widths = [2008, 1659, 1663, 1379, 1687, 1969, 2035, 2389]
    _set_fixed_table_layout(table)
    _set_table_grid(table, widths)
    _set_table_width(table, sum(widths))
    _set_table_borders(table)

    _set_row_height(table.rows[0], 0.5465)
    _set_row_height(table.rows[1], 0.3035)

    headers = [
        "Total Weighed (X)",
        "Warned & Charged (Y)",
        "Charged & Prohibited (Z)",
        "Warned (A)",
        "Cases Cleared in Court (B)",
        "  Transgressions (L)",
        "Exempted permit (E)",
        "Manually Weighed (M)",
    ]
    values = _summary_values(records, session)
    body = [
        values["total_weighed"],
        values["warned_and_charged"],
        values["charged"],
        values["warned"],
        values["cases_cleared"],
        values["transgressions"],
        values["exempted_permit"],
        values["manually_weighed"],
    ]
    for col, header in enumerate(headers):
        _set_cell_text(table.cell(0, col), header, size=11, bold=True)
        _set_cell_width(table.cell(0, col), widths[col])
        _set_cell_text(table.cell(1, col), body[col], font_name="Calibri", size=11)
        _set_cell_width(table.cell(1, col), widths[col])


def _add_nil_table(doc: Document, heading: str, headers: list[str], widths: list[int]) -> None:
    _add_heading(doc, heading)
    table = doc.add_table(rows=2, cols=len(headers))
    _set_fixed_table_layout(table)
    _set_table_grid(table, widths)
    _set_table_width(table, sum(widths))
    _set_table_borders(table)

    if heading.upper() == "TRANSGRESSION":
        _set_row_height(table.rows[0], 0.1639)
        _set_row_height(table.rows[1], 0.2569)
    elif "REPORT" in heading.upper():
        _set_row_height(table.rows[0], 0.0868)
        _set_row_height(table.rows[1], 0.2660)
    else:
        # Number plates removed
        _set_row_height(table.rows[0], 0.0840)
        _set_row_height(table.rows[1], 0.2576)

    for col, header in enumerate(headers):
        _set_cell_text(table.cell(0, col), header, size=11, bold=True)
        _set_cell_width(table.cell(0, col), widths[col])
        _set_cell_text(table.cell(1, col), "NIL" if col == 0 else "", font_name="Calibri", size=11)
        _set_cell_width(table.cell(1, col), widths[col])


def _vehicle_row_values(record, session, report_date) -> list[str]:
    gvw_excess = record["gvw_difference_kg"] if record["gvw_difference_kg"] > 0 else 0
    axle_excess = record["excess_kg"] if record["excess_kg"] > 0 else 0
    danka_staff = _manual_value(session, "danka_staff", "computer_operator")
    police_officers = _manual_value(session, "police_officers", "police")

    return [
        _vehicle_display_date(report_date),
        _upper(record["registration"]),
        _upper(record["transporter"]),
        _upper(record["axle"]),
        _format_number(gvw_excess),
        _format_number(axle_excess),
        _upper(record["origin"]),
        _upper(record["destination"]),
        _upper(record["cargo"]),
        _format_location_people(danka_staff).rstrip("."),
        _format_location_people(police_officers).rstrip("."),
        _upper(record["remarks"]),
    ]


def _add_vehicle_table(
    doc: Document,
    heading: str,
    records: pd.DataFrame,
    session,
    report_date,
) -> None:
    _add_heading(doc, heading)
    rows = max(1, len(records)) + 2
    table = doc.add_table(rows=rows, cols=12)
    table.autofit = False
    table.allow_autofit = False
    _set_fixed_table_layout(table)
    _set_table_grid(table, VEHICLE_TABLE_WIDTHS)
    _set_table_width(table, sum(VEHICLE_TABLE_WIDTHS))
    _set_table_borders(table, size="6")

    headers = [
        "DATE WEIGHED",
        "VEHICLE REG",
        "TRANSPORTER",
        "CONFIG.",
        "EXCESS   WEIGHT",
        "EXCESS   WEIGHT",
        "ORIGIN",
        "DESTIN.",
        "CARGO",
        "COMPUTER OPERATOR (DANKA STAFF)",
        "OFFICER",
        "REMARKS",
    ]
    subheaders = ["", "", "", "", "GVW", "AXLE", "", "", "", "", "", ""]
    for col, header in enumerate(headers):
        _set_cell_text(table.cell(0, col), header, size=8, bold=True)
        _set_cell_width(table.cell(0, col), VEHICLE_TABLE_WIDTHS[col])
        _set_cell_text(table.cell(1, col), subheaders[col], size=8, bold=True)
        _set_cell_width(table.cell(1, col), VEHICLE_TABLE_WIDTHS[col])

    cell_0_4 = table.cell(0, 4)
    cell_0_5 = table.cell(0, 5)
    cell_0_4.merge(cell_0_5)
    _set_cell_text(cell_0_4, "EXCESS   WEIGHT", size=8, bold=True)

    # Table 6 (full vehicle list) vs Table 7 (charged >2T) have different header row heights
    if "CHARGED" in heading.upper():
        _set_row_height(table.rows[0], 0.7729)
        _set_row_height(table.rows[1], 0.2264)
    else:
        _set_row_height(table.rows[0], 0.5507)
        _set_row_height(table.rows[1], 0.1826)

    if records.empty:
        for col in range(12):
            _set_cell_text(table.cell(2, col), "NIL" if col == 0 else "", font_name="Calibri", size=11)
            _set_cell_width(table.cell(2, col), VEHICLE_TABLE_WIDTHS[col])
    else:
        for row_index, (_, record) in enumerate(records.iterrows(), start=2):
            row_values = _vehicle_row_values(record, session, report_date)
            for col, value in enumerate(row_values):
                align = WD_ALIGN_PARAGRAPH.RIGHT if col in (4, 5) else WD_ALIGN_PARAGRAPH.LEFT
                _set_cell_text(table.cell(row_index, col), value, font_name="Calibri", size=11, align=align)
                _set_cell_width(table.cell(row_index, col), VEHICLE_TABLE_WIDTHS[col])


def _add_mileage_table(doc: Document, session) -> None:
    _add_heading(doc, "LOCATION REPORT", size=12)
    table = doc.add_table(rows=4, cols=4)
    widths = [4235, 3980, 2625, 3560]
    _set_fixed_table_layout(table)
    _set_table_grid(table, widths)
    _set_table_width(table, sum(widths))
    _set_table_borders(table, size="12")

    _set_row_height(table.rows[0], 0.4375)
    _set_row_height(table.rows[1], 0.2500)
    _set_row_height(table.rows[2], 0.2188)
    _set_row_height(table.rows[3], 0.2188)

    start = _manual_value(session, "mileage_start")
    end = _manual_value(session, "mileage_end")
    kms = ""
    start_number = _plain_int(start)
    end_number = _plain_int(end)
    if start_number is not None and end_number is not None:
        kms = _format_number(end_number - start_number)

    headers = ["MILEAGE START", "MILEAGE END", "KMS", "MOBILE VEHICLE"]
    values = [
        _format_number(start_number) if start_number is not None else "",
        _format_number(end_number) if end_number is not None else "",
        kms,
        _upper(_manual_value(session, "mobile_vehicle", "vehicle_used", "vehicle")),
    ]

    for col, header in enumerate(headers):
        _set_cell_text(
            table.cell(0, col),
            header.upper(),
            font_name="Arial",
            size=12,
            bold=True,
            valign=WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        )
        _set_cell_width(table.cell(0, col), widths[col])
        _set_cell_text(
            table.cell(1, col),
            values[col],
            font_name="Calibri",
            size=14,
            valign=WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        )
        _set_cell_width(table.cell(1, col), widths[col])

    for col in range(4):
        _set_cell_text(
            table.cell(2, col),
            "",
            font_name="Calibri",
            size=14,
            valign=WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        )
        _set_cell_width(table.cell(2, col), widths[col])
    _set_cell_text(
        table.cell(3, 2),
        f"{kms} KMS" if kms else "",
        font_name="Calibri",
        size=14,
        bold=True,
        valign=WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    )


def _split_location_people(value: Any) -> list[str]:
    if value in (None, ""):
        return []

    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    parts = re.split(r"\s*/\s*|\n+", text)
    return [part.strip().strip(".") for part in parts if part.strip().strip(".")]


def _format_location_people(value: Any) -> str:
    names = [_upper(name) for name in _split_location_people(value)]
    if not names:
        return "."
    return f"{' / '.join(names)}."


def _format_route(value: Any) -> str:
    route = _upper(value).strip(".")
    return f"{route}." if route else "."


def _add_location_notes(doc: Document, session) -> None:
    route = _manual_value(session, "route")
    danka_staff = _manual_value(session, "danka_staff", "computer_operator")
    police = _manual_value(session, "police_officers", "police")

    _add_mixed_line(
        doc,
        "ACTUAL ROUTE: ",
        _format_route(route),
        space_before=6,
        space_after=6,
    )
    _add_mixed_line(
        doc,
        "DANKA PERSONNEL : ",
        _format_location_people(danka_staff),
        space_after=6,
    )
    _add_mixed_line(
        doc,
        "POLICE OFFICERS : ",
        _format_location_people(police),
    )


def _style_footer_run(run) -> None:
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")


def _add_footer_field(paragraph, field_name) -> None:
    run = paragraph.add_run()
    _style_footer_run(run)

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field_name} "

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")

    field_text = OxmlElement("w:t")
    field_text.text = "1"

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_separate)
    run._r.append(field_text)
    run._r.append(field_end)


def _add_mobile_footer(section, report_date, session) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    printable_width = (
        section.page_width.inches
        - section.left_margin.inches
        - section.right_margin.inches
    )
    tab_stops = paragraph.paragraph_format.tab_stops
    report_label_tab = min(3.35, max(3.0, printable_width * 0.34))
    tab_stops.add_tab_stop(Inches(report_label_tab), WD_TAB_ALIGNMENT.LEFT)
    tab_stops.add_tab_stop(Inches(printable_width), WD_TAB_ALIGNMENT.RIGHT)

    station = _mobile_station_label(session)
    report_number = _mobile_report_number(session)

    date_obj = _date_value(report_date)
    day = date_obj.day
    suffix = "TH" if 11 <= day <= 13 else {1: "ST", 2: "ND", 3: "RD"}.get(day % 10, "TH")
    month_year = date_obj.strftime("%B, %Y").upper()

    reference = paragraph.add_run("KeNHA/WB/MTCE/4339/2025")
    _style_footer_run(reference)

    first_tab = paragraph.add_run("\t")
    _style_footer_run(first_tab)

    report_label = paragraph.add_run(f"{station} MOBILE REPORT {report_number} {day}")
    _style_footer_run(report_label)

    suffix_label = paragraph.add_run(suffix)
    _style_footer_run(suffix_label)
    suffix_label.font.superscript = True

    month_year_label = paragraph.add_run(f" {month_year}")
    _style_footer_run(month_year_label)

    second_tab = paragraph.add_run("\tPage ")
    _style_footer_run(second_tab)

    _add_footer_field(paragraph, "PAGE")
    run_end = paragraph.add_run(" of ")
    _style_footer_run(run_end)
    _add_footer_field(paragraph, "NUMPAGES")


def _add_mobile_title_header(section, session) -> None:
    header = section.header
    for paragraph in list(header.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for table in list(header.tables):
        table._element.getparent().remove(table._element)

    printable_width = (
        section.page_width.inches
        - section.left_margin.inches
        - section.right_margin.inches
    )
    table = header.add_table(rows=1, cols=3, width=Inches(printable_width))
    table.allow_autofit = False
    _set_table_width(table, int(printable_width * 1440))
    _remove_table_borders(table)

    logo_width = 1.15
    right_spacer_width = 0.1
    center_width = max(printable_width - logo_width - right_spacer_width, 2.0)
    column_widths = [logo_width, center_width, right_spacer_width]
    column_width_twips = [int(width * 1440) for width in column_widths]
    _set_table_grid(table, column_width_twips)
    for col, width in enumerate(column_width_twips):
        _set_cell_width(table.cell(0, col), width)

    station = _mobile_station_label(session)
    report_number = _mobile_report_number(session)

    logo_cell = table.cell(0, 0)
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_paragraph.paragraph_format.space_before = Pt(0)
    logo_paragraph.paragraph_format.space_after = Pt(0)
    if LOGO_PATH.exists():
        logo_paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.0))

    title_cell = table.cell(0, 1)
    title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_no_wrap(title_cell)
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(0)

    run = title_paragraph.add_run(f"{station} MOBILE DAILY REPORT {report_number}")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def build_mobile_word_report(session) -> io.BytesIO:
    if (
        "mobile_report" not in session.dataframes
        or session.sections.get("mobile_report", {}).get("status") != "ready"
    ):
        raise ValueError("Mobile report data is not ready")

    records = session.dataframes["mobile_report"].copy()
    if records.empty:
        raise ValueError("Mobile report data is empty")

    records["date_time"] = pd.to_datetime(records["date_time"], errors="coerce")
    records = records.loc[records["date_time"].notna()].copy()
    if records.empty:
        raise ValueError("Mobile report data has no valid Date Time values")

    report_date = _date_value(records["date_time"].iloc[0])
    doc = Document()
    enable_field_updates(doc)
    
    section_0 = doc.sections[0]
    _set_section_portrait(section_0)
    _add_mobile_title_header(section_0, session)
    _add_mobile_footer(section_0, report_date, session)

    _add_daily_hour_statistics(doc, records, report_date)
    _add_prepared_approved(doc, session)

    _add_initial_landscape_section(doc, report_date, session)
    _add_daily_hourly_data(doc, records, session)

    _add_continuous_landscape_section(doc)
    _add_daily_summary(doc, records, session)

    _add_continuous_landscape_section(doc)
    _add_nil_table(
        doc,
        "4.0 TRANSGRESSION",
        [
            "Date",
            "Time received",
            "Truck No",
            "Sending station",
            "CI reported to",
            "Action 1",
            "Action 2",
            "Attach evidence",
            "Caught Or Not Caught",
            "Next Weighbridge Report Sent",
            "Next Weighbridge",
        ],
        [921, 1369, 1012, 1337, 1369, 1263, 1115, 1428, 1204, 1902, 1902],
    )

    transgression_headers = [
        "Date",
        "Time",
        "Reg. No.",
        "Axle Config",
        "Transporter",
        "Computer Operator",
        "Police In charge",
        "Action Taken",
        "Caught or not",
        "Next WB report sent",
        "Next WB",
        "Next WB Recepient",
    ]

    _add_continuous_landscape_section(doc)
    _add_nil_table(
        doc,
        "I. TRANSGRESSION REPORT",
        transgression_headers,
        [1044, 949, 949, 1044, 1614, 1424, 1543, 1206, 1276, 1482, 864, 1447],
    )

    _add_continuous_landscape_section(doc)
    _add_nil_table(
        doc,
        "5.0 NUMBER plates removed",
        transgression_headers,
        [1046, 951, 951, 1046, 1617, 1427, 1546, 1209, 1279, 1485, 866, 1450],
    )

    _add_continuous_landscape_section(doc)
    _add_vehicle_table(doc, "6.0 DETAILS OF VEHICLES", records, session, report_date)

    charged_over_two = records.loc[
        (records["gvw_difference_kg"] > 2000)
        & (records["is_gvw_axle_charge"] | records["is_dimension_charge"])
    ].copy()
    _add_continuous_landscape_section(doc)
    _add_vehicle_table(
        doc,
        "7.0 CHARGED ABOVE 2 TONNES ON GVW",
        charged_over_two,
        session,
        report_date,
    )

    _add_continuous_landscape_section(doc)
    _add_mileage_table(doc, session)

    _add_continuous_landscape_section(doc)
    _add_location_notes(doc, session)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
