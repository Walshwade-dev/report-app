from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.services.daily_summary_processor import daily_summary_rows


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table):
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)

    tbl_layout.set(qn("w:type"), "fixed")


def set_table_grid(table, widths):
    tbl = table._tbl

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    tbl_grid = OxmlElement("w:tblGrid")

    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    tbl.insert(0, tbl_grid)


def style_cell(
    cell,
    font_size=7,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for paragraph in cell.paragraphs:
        paragraph.alignment = align

        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.bold = bold

            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)

            r_fonts.set(qn("w:ascii"), "Arial")
            r_fonts.set(qn("w:hAnsi"), "Arial")


def add_daily_summary_section(doc: Document, summary: dict):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run("4. DAILY SUMMARY")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    headers_1 = [
        "Weighed by\nHSWIM\n(Q)",
        "Weighed\nScale\ntotal\n(N)=D\n+S",
        "Manually\nWeighed\n(M)",
        "Total\nweighed\n(X)",
        "Total\nTraffic\n(T)",
        "Total\nOverload\n(Y)\nA+Z+G\n+R",
        "Warned\n(A)",
        "Charged\n&Prohibited\n(Z)",
        "Special\nrelease\n(G)",
        "Vehicles\nCharged\nbut\nRedistributed\n(R)",
        "Impounded\n& prohibited\n(P)\nZ+R+G",
        "Cases\ncleared\nin\nCourt(\nB)",
        "Transgressions",
        "Exemption permits",
        "",
        "",
    ]

    headers_2 = [
        "(Q=H-C)",
        "(N)",
        "(M)",
        "(X)=(S+M)",
        "(T)=(Q+X+K+E)",
        "(Y)",
        "(A)",
        "(Z)",
        "(G)",
        "(R)",
        "(P)",
        "(B)",
        "(L)",
        "Not\nweighed\n(E)",
        "Weighed\n(F)",
        "Total",
    ]

    values = [
        summary["weighed_by_hswim_q"],
        summary["weighed_scale_total_n"],
        summary["manually_weighed_m"],
        summary["total_weighed_x"],
        summary["total_traffic_t"],
        summary["total_overload_y"],
        summary["warned_a"],
        summary["charged_prohibited_z"],
        summary["special_release_g"],
        summary["vehicles_charged_but_redistributed_r"],
        summary["impounded_prohibited_p"],
        summary["cases_cleared_in_court_b"],
        summary["transgressions_l"],
        summary["exemption_permits_not_weighed_e"],
        summary["exemption_permits_weighed_f"],
        summary["exemption_permits_total"],
    ]

    table = doc.add_table(rows=3, cols=16)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    widths = [
        918, 828, 864, 1170, 971, 889, 750, 817,
        730, 1162, 1007, 874, 900, 1069, 828, 803
    ]

    set_fixed_table_layout(table)
    set_table_grid(table, widths)

    for i, text in enumerate(headers_1):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_width(cell, widths[i])
        style_cell(cell, font_size=10, bold=True)

    for i, text in enumerate(headers_2):
        cell = table.rows[1].cells[i]
        cell.text = text
        set_cell_width(cell, widths[i])
        style_cell(cell, font_size=10, bold=True)

    for i, value in enumerate(values):
        cell = table.rows[2].cells[i]
        cell.text = f"{int(value):,}"
        set_cell_width(cell, widths[i])
        style_cell(cell, font_size=11)

    # Merge the top "Exemption permits" header across E, F, Total
    exemption_cell = table.rows[0].cells[13]
    exemption_cell.merge(table.rows[0].cells[15])

    for row in table.rows:
        row.height = Pt(25.9)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
