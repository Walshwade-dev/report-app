import io

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from app.services.report_layout import (
    FONT_NAME,
    SECTION_TITLE_SIZE,
    SUBHEADING_SIZE,
    TABLE_HEADER_SIZE,
    TABLE_BODY_SIZE,
)

TABLE_WIDTHS = [1110, 960, 660, 1060, 1050]
COMPACT_TABLE_WIDTHS = TABLE_WIDTHS
# Container: left column holds the data table, right column holds the chart.
# Total printable width = 15,000 twips (11" - 1" margins at 1440 twips/inch).
# Left matching the exact compact table width (4840 twips), right takes the remaining (10160 twips).
SECTION_TWO_CONTAINER_WIDTHS = [4840, 10160]
# Chart dimensions that fit within the printable page height.
CHART_WIDTH_INCHES = 6.61
CHART_HEIGHT_INCHES = 6.57


def set_row_height(row, height):
    row.height = Inches(height)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_row_height_points(row, points):
    row.height = Pt(points)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))

    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table):
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))

    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)

    tbl_layout.set(qn("w:type"), "fixed")


def set_table_grid(table, widths):
    tbl = table._tbl
    existing_grid = tbl.find(qn("w:tblGrid"))

    if existing_grid is not None:
        tbl.remove(existing_grid)

    tbl_grid = OxmlElement("w:tblGrid")

    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    tbl.insert(0, tbl_grid)


def set_table_borders(table, size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))

        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def set_cell_borders(cell, value="single", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))

        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)

        border.set(qn("w:val"), value)
        if value != "nil":
            border.set(qn("w:sz"), size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))

        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)

        border.set(qn("w:val"), "nil")


def set_table_cell_margins(table, top=0, bottom=0, left=0, right=0):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is not None:
        tbl_pr.remove(tbl_cell_mar)
    tbl_cell_mar = OxmlElement("w:tblCellMar")
    for margin, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{margin}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tbl_cell_mar.append(node)
    tbl_pr.append(tbl_cell_mar)


def set_table_indent(table, width):
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))

    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)

    tbl_ind.set(qn("w:w"), str(width))
    tbl_ind.set(qn("w:type"), "dxa")


def apply_widths_to_all_cells(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])


def remove_empty_cell_paragraphs(cell):
    for paragraph in cell.paragraphs:
        if paragraph.text:
            continue

        paragraph._element.getparent().remove(paragraph._element)


def set_cell_text(cell, text):
    cell.text = str(text)


def format_count(value):
    if pd.isna(value):
        return ""

    return f"{int(value):,}"


def build_daily_hour_chart_data(daily_df: pd.DataFrame) -> pd.DataFrame:
    """
    Build data used for Daily Hourly Data table and chart.

    N = D + S
    M = M
    Q = H - C
    X = N + M
    """

    df = daily_df[daily_df["DATE"].astype(str).str.lower() != "totals"].copy()

    chart_df = pd.DataFrame()

    chart_df["TIME"] = df["TIME"]
    chart_df["N"] = df["D"] + df["S"]
    chart_df["M"] = df["M"]
    chart_df["Q"] = df["Q"]
    chart_df["X"] = chart_df["N"] + chart_df["M"]

    return chart_df


def create_daily_hour_chart_image(chart_df, width=6.61, height=6.57):
    buffer = io.BytesIO()

    fig = plt.figure(figsize=(width, height), dpi=150)
    ax = fig.add_axes([0.08, 0.18, 0.89, 0.72])
    fig.patch.set_facecolor("white")
    fig.patch.set_edgecolor("#D9D9D9")
    fig.patch.set_linewidth(1.0)

    max_val = max(
        chart_df["N"].max() if not chart_df["N"].empty else 0,
        chart_df["M"].max() if not chart_df["M"].empty else 0,
        chart_df["Q"].max() if not chart_df["Q"].empty else 0,
        chart_df["X"].max() if not chart_df["X"].empty else 0
    )
    max_val = max(max_val, 300)
    ylim_top = ((int(max_val) + 49) // 50) * 50

    ax.set_ylim(-ylim_top * 0.04, ylim_top * 1.06)
    ax.set_yticks(range(0, ylim_top + 1, 50))
    ax.plot(chart_df["TIME"], chart_df["N"], label="N=(D+S)", color="#4472C4", linewidth=2.2)
    ax.plot(chart_df["TIME"], chart_df["M"], label="(M)", color="#ED7D31", linewidth=2.2)
    ax.plot(chart_df["TIME"], chart_df["Q"], label="Q= H-C", color="#A5A5A5", linewidth=2.6)
    ax.plot(chart_df["TIME"], chart_df["X"], label="X= (D+S+M)", color="#FFC000", linewidth=2.6)

    ax.set_title(
        "Graph on Trucks Weighed per Hour",
        fontsize=14,
        fontweight="bold",
        color="#595959",
        pad=12,
    )
    ax.tick_params(axis="x", labelrotation=48, labelsize=7, colors="#595959")
    ax.tick_params(axis="y", labelsize=8, colors="#595959")
    ax.grid(axis="y", color="#D9D9D9", linewidth=0.9)
    ax.grid(axis="x", visible=False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#BFBFBF")
    ax.spines["bottom"].set_color("#BFBFBF")
    ax.legend(
        fontsize=7,
        loc="center",
        bbox_to_anchor=(0.5, 0.08),
        bbox_transform=fig.transFigure,
        ncol=4,
        frameon=False,
        handlelength=2.4,
        handletextpad=0.3,
        columnspacing=1.1,
    )

    fig.savefig(buffer, format="png", dpi=150)
    plt.close()

    buffer.seek(0)
    return buffer


def style_table_cell(
    cell,
    font_size=10,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.bold = bold

            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)

            r_fonts.set(qn("w:ascii"), "Arial")
            r_fonts.set(qn("w:hAnsi"), "Arial")


def add_vertical_spacer(doc, points):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(points)

    run = paragraph.add_run("")
    run.font.size = Pt(1)


def add_daily_hour_chart_section(doc: Document, daily_df, is_preview: bool = False):
    chart_df = build_daily_hour_chart_data(daily_df)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(4)
    # Start section 2 on its own page by setting page_break_before.
    # This is cleaner than an external doc.add_page_break() call because
    # it keeps the heading and the table that follows it atomic.
    heading.paragraph_format.page_break_before = True
    run = heading.add_run("2. DAILY HOURLY DATA")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    container_table = doc.add_table(rows=1, cols=2)
    container_table.autofit = False
    container_table.allow_autofit = False
    container_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_fixed_table_layout(container_table)
    set_table_grid(container_table, SECTION_TWO_CONTAINER_WIDTHS)
    remove_table_borders(container_table)
    set_table_cell_margins(container_table, top=0, bottom=0, left=0, right=0)

    left_cell = container_table.rows[0].cells[0]
    right_cell = container_table.rows[0].cells[1]
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_width(left_cell, SECTION_TWO_CONTAINER_WIDTHS[0])
    set_cell_width(right_cell, SECTION_TWO_CONTAINER_WIDTHS[1])

    small_table = left_cell.add_table(rows=27, cols=5)
    small_table.autofit = False
    small_table.allow_autofit = False
    small_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    set_fixed_table_layout(small_table)
    set_table_grid(small_table, COMPACT_TABLE_WIDTHS)
    set_table_indent(small_table, 0)
    set_table_borders(small_table, size="6")

    if is_preview:
        header_height = 24.0
        formula_height = 24.0
        row_height = 14.0
        chart_w = 6.61
        chart_h = 5.0
    else:
        header_height = 31.68
        formula_height = 31.68
        row_height = 16.56
        chart_w = CHART_WIDTH_INCHES
        chart_h = CHART_HEIGHT_INCHES

    header_row = small_table.rows[0]
    formula_row = small_table.rows[1]

    set_row_height_points(header_row, header_height)
    set_row_height_points(formula_row, formula_height)

    time_cell = header_row.cells[0].merge(formula_row.cells[0])

    headers = [
        "Time",
        "Multideck\nweighed",
        "Manually",
        "HSWIM\nCLEARED",
        "Total\nweighed",
    ]

    for i, header in enumerate(headers):
        cell = time_cell if i == 0 else header_row.cells[i]
        set_cell_text(cell, header)
        set_cell_width(cell, COMPACT_TABLE_WIDTHS[i])
        style_table_cell(cell, font_size=8, bold=True)
        set_cell_borders(cell, size="6")

    set_cell_borders(header_row.cells[0], size="6")
    set_cell_borders(formula_row.cells[0], size="6")

    formulas = ["N=(D+S)", "(M)", "Q = H-C", "X= (N+M)"]

    for i, formula in enumerate(formulas, start=1):
        cell = formula_row.cells[i]
        set_cell_text(cell, formula)
        set_cell_width(cell, COMPACT_TABLE_WIDTHS[i])
        style_table_cell(cell, font_size=8, bold=True)
        set_cell_borders(cell, size="6")

    for row_index, (_, row) in enumerate(chart_df.iterrows(), start=2):
        row_obj = small_table.rows[row_index]
        set_row_height_points(row_obj, row_height)

        values = [row["TIME"], row["N"], row["M"], row["Q"], row["X"]]

        for i, value in enumerate(values):
            cell = row_obj.cells[i]
            set_cell_text(cell, value if i == 0 else format_count(value))
            set_cell_width(cell, COMPACT_TABLE_WIDTHS[i])

            style_table_cell(
                cell,
                font_size=8,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            set_cell_borders(cell, size="6")

    totals = small_table.rows[-1]
    set_row_height_points(totals, row_height)

    total_values = [
        "Total",
        chart_df["N"].sum(),
        chart_df["M"].sum(),
        chart_df["Q"].sum(),
        chart_df["X"].sum(),
    ]

    for i, value in enumerate(total_values):
        cell = totals.cells[i]
        set_cell_text(cell, value if i == 0 else format_count(value))
        set_cell_width(cell, COMPACT_TABLE_WIDTHS[i])
        style_table_cell(
            cell,
            font_size=8,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT,
        )
        set_cell_borders(cell, size="6")

    chart_image = create_daily_hour_chart_image(chart_df, width=chart_w, height=chart_h)

    paragraph = right_cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(chart_image, width=Inches(chart_w), height=Inches(chart_h))

    for cell in [left_cell, right_cell]:
        for p in list(cell.paragraphs):
            if not p.text and len(p.runs) == 0:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(1)
