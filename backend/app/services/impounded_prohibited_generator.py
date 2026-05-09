from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
)
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.report_layout import A4_TABLE_WIDTH_TWIPS, apply_standard_layout

import io
import pandas as pd


HEADER_LABELS = {
    "Date Weighed/Prohibited": "Date\nWeighed/\nProhibited",
    "Axle Config": "Axle\nConfig",
    "GVW Over load": "GVW\nOver\nload",
    "Computer Operator": "Computer\nOperator",
}


COLUMN_RATIOS = {
    "Date Weighed/Prohibited": 1.3,
    "Transporter": 1.6,
    "Cargo": 1.6,
    "Source": 1.3,
    "Destination": 1.3,
    "ProhibitionOrder": 2.2,
    "Prosecutor": 1.5,
    "Computer Operator": 1.7,
}


def get_column_widths(columns, total_width=A4_TABLE_WIDTH_TWIPS):
    sample_widths = [1279, 1254, 1113, 835, 1669, 1116, 1255, 976, 696, 2229, 1206, 1492]
    if len(columns) == len(sample_widths):
        return {col: sample_widths[index] for index, col in enumerate(columns)}

    total_ratio = sum(COLUMN_RATIOS.get(col, 1) for col in columns)
    base_width = total_width / total_ratio

    return {
        col: int(base_width * COLUMN_RATIOS.get(col, 1))
        for col in columns
    }


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table):
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)

    tbl_layout.set(qn("w:type"), "fixed")


def set_table_grid(table, columns, widths):
    tbl = table._tbl

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    tbl_grid = OxmlElement("w:tblGrid")

    for col in columns:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(widths[col]))
        tbl_grid.append(grid_col)

    tbl.insert(0, tbl_grid)


def apply_widths_to_all_cells(table, columns, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[columns[i]])


def style_cell(
    cell,
    font_size=7,
    bold=False,
    valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
    align=WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.vertical_alignment = valign

    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.bold = bold


def add_impounded_prohibited_section(doc: Document, df: pd.DataFrame):

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run("6. IMPOUNDED & PROHIBITED")
    run.bold = True
    run.underline = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    columns = list(df.columns)
    widths = get_column_widths(columns)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    set_fixed_table_layout(table)
    set_table_grid(table, columns, widths)

    header_row = table.rows[0]
    header_row.height = Pt(1.2)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for i, col in enumerate(columns):
        cell = header_row.cells[i]
        cell.text = HEADER_LABELS.get(col, col)

        style_cell(
            cell,
            font_size=8,
            bold=True,
        )

    for _, row in df.iterrows():
        row_obj = table.add_row()
        row_obj.height = Pt(56.7)
        row_obj.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        for i, value in enumerate(row):
            cell = row_obj.cells[i]
            cell.text = "" if pd.isna(value) else str(value).upper()

            style_cell(
                cell,
                font_size=10,
            )
            
        apply_widths_to_all_cells(table, columns, widths)

def generate_impounded_prohibited_report(
    df: pd.DataFrame,
    report_date: str,
    station: str,
    bound: str,
) -> io.BytesIO:
    
    doc = Document()

    apply_standard_layout(
        doc,
        report_date=report_date,
        station=station,
        bound=bound,
    )

    add_impounded_prohibited_section(doc, df)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
