from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
)
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.services.report_layout import (
    FONT_NAME,
    SECTION_TITLE_SIZE,
    SUBHEADING_SIZE,
    TABLE_HEADER_SIZE,
    TABLE_BODY_SIZE,
)
import pandas as pd

from app.services.report_layout import A4_TABLE_WIDTH_TWIPS


SAMPLE_COLUMN_WIDTHS = [
    1096,
    988,
    1078,
    899,
    899,
    899,
    934,
    965,
    810,
    810,
    1170,
    900,
    1080,
    900,
    884,
    1096,
]


HEADER_LABELS = {
    "DATE": "DATE",
    "TIME": "TIME",
    "D": "MULTIDECK\nSCALE\n(D)",
    "S": "WEIGHED\nSAW\n(S)",
    "M": "MANUAL\n(M)",
    "H": "HSWIM\nTOTAL\n(H)",
    "Q": "HSWIM -\nCLEARED\nQ = H-C",
    "X": "TOTAL\nWEIGHED\nX=(D+S+M)",
    "C": "CALLED\nIN\n(C)",
    "Y": "TOTAL\nOVERLOADED\nY=(A+Z+G+R)",
    "P": "IMPOUNDED &\nPROHIBITED\nP=(Z+R)",
    "A": "WARNED\nTRUCKS\n(A)",
    "Z": "CHARGED &\nPROHIBITED\n(Z)",
    "G": "SPECIAL\nRELEASE\n(G)",
    "R": "REDISTRIBUTED\n(R)",
    "E": "EXEMPTION\nPERMITS NOT\nWEIGHED\n(E)",
}


COLUMN_RATIOS = {
    "DATE": 1.2,
    "TIME": 1.1,
    "Y": 1.4,
    "P": 1.5,
    "E": 1.5,
}


def get_column_widths(columns, total_width=A4_TABLE_WIDTH_TWIPS):
    if len(columns) == len(SAMPLE_COLUMN_WIDTHS):
        return {col: SAMPLE_COLUMN_WIDTHS[index] for index, col in enumerate(columns)}

    total_ratio = sum(COLUMN_RATIOS.get(col, 1) for col in columns)
    base_width = total_width / total_ratio

    return {
        col: int(base_width * COLUMN_RATIOS.get(col, 1))
        for col in columns
    }


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table):
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)

    tbl_layout.set(qn("w:type"), "fixed")


def set_table_grid(table, columns, widths):
    tbl = table._tbl

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    tbl_grid = OxmlElement("w:tblGrid")

    for col in columns:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(widths[col]))
        tbl_grid.append(grid_col)

    tbl.insert(0, tbl_grid)


def apply_widths_to_all_cells(table, columns, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[columns[i]])


def style_cell(
    cell,
    font_size=10,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for paragraph in cell.paragraphs:
        paragraph.alignment = align

        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.bold = bold

            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)

            r_fonts.set(qn("w:ascii"), "Arial")
            r_fonts.set(qn("w:hAnsi"), "Arial")


def add_daily_hour_statistics_section(doc: Document, daily_df: pd.DataFrame):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run("1. DAILY AND HOURLY STATISTICS")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    columns = list(daily_df.columns)
    widths = get_column_widths(columns)

    table = doc.add_table(rows=3, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    set_fixed_table_layout(table)
    set_table_grid(table, columns, widths)

    row1 = table.rows[0]
    row2 = table.rows[1]
    row3 = table.rows[2]

    row1.height = Pt(14.5)
    row2.height = Pt(31.0)
    row3.height = Pt(24.0)

    row1.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row2.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row3.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # vertical merges for DATE/TIME
    row1.cells[0].merge(row3.cells[0])
    row1.cells[1].merge(row3.cells[1])

    row1.cells[0].text = "DATE"
    row1.cells[1].text = "TIME"

    # grouped header
    trucks_cell = row1.cells[2]
    for i in range(3, 8):
        trucks_cell = trucks_cell.merge(row1.cells[i])
    trucks_cell.text = "TRUCKS WEIGHED"

    # other top headers
    top_headers = {
        8: "CALLED\nIN",
        9: "TOTAL\nOVERLOADED",
        10: "IMPOUNDED\n&\nPROHIBITED",
        11: "WARNED\nTRUCKS",
        12: "CHARGED &\nPROHIBITED",
        13: "SPECIAL\nRELEASE",
        14: "REDISTRI\nBUTED",
        15: "EXEMPTION\nPERMITS\nNOT\nWEIGHED",
    }

    for idx, label in top_headers.items():
        row1.cells[idx].merge(row2.cells[idx])
        row1.cells[idx].text = label

    # second-row labels under TRUCKS WEIGHED
    row2_labels = {
        2: "MULTIDECK\nSCALE",
        3: "WEIGHED\nSAW",
        4: "MANUAL",
        5: "HSWIM\nTOTAL",
        6: "HSWIM -\nCLEARED",
        7: "TOTAL\nWEIGHED",
    }

    for idx, label in row2_labels.items():
        row2.cells[idx].text = label

    # third-row formulas
    row3_labels = {
    2: "(D)",
    3: "(S)",
    4: "(M)",
    5: "(H)",
    6: "Q = H-C",
    7: "X=(D\n+M+S)",
    8: "(C)",
    9: "(Y)=(A+Z+G+R)",
    10: "(P)=(Z+R)",
    11: "(A)",
    12: "(Z)",
    13: "(G)",
    14: "(R)",
    15: "(E)",
}

    for idx, label in row3_labels.items():
        row3.cells[idx].text = label

    # style headers
    for row in [row1, row2, row3]:
        for cell in row.cells:
            style_cell(cell, font_size=7, bold=True)

    # data rows INCLUDING totals row
    for _, row in daily_df.iterrows():
        row_obj = table.add_row()
        row_obj.height = Pt(13.8)
        row_obj.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


        is_total_row = str(row.get("DATE", "")).strip().lower() == "totals"

        for i, value in enumerate(row):
            cell = row_obj.cells[i]

            if pd.isna(value):
                text = ""
            elif isinstance(value, float) and value.is_integer():
                text = str(int(value))
            else:
                text = str(value)

            cell.text = text
            is_date_or_time = i in {0, 1}
            font_size = 10 if is_date_or_time else 11
            style_cell(cell, font_size=font_size, bold=is_total_row or is_date_or_time)

    apply_widths_to_all_cells(table, columns, widths)
    
