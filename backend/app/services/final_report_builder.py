from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import pandas as pd

from app.services.report_layout import FONT_NAME, REPORT_TITLE_SIZE, apply_standard_layout
from app.services.report_context import ReportContext
from app.services.report_session_metrics import get_daily_hour_totals_from_dataframe
from app.services.wideload_generator import add_wideload_section
from app.services.impounded_prohibited_generator import add_impounded_prohibited_section
from app.services.overloaded_summary import count_valid_permit_vehicles
from app.services.daily_hour_generator import add_daily_hour_statistics_section
from app.services.daily_hour_chart_generator import add_daily_hour_chart_section
from app.services.traffic_census_generator import add_traffic_census_section
from app.services.daily_summary_generator import add_daily_summary_section
from app.services.daily_summary_processor import build_daily_summary
from app.services.transgressions_generator import add_transgressions_section


def add_report_title(doc: Document, station: str, bound: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(
        f"{station.upper()} WEIGHBRIDGE {bound.upper()} DAILY REPORT"
    )
    run.bold = True
    run.underline = True
    run.font.name = FONT_NAME
    run.font.size = Pt(REPORT_TITLE_SIZE)


def add_prepared_confirmed_lines(
    doc: Document,
    prepared_by: str | None = None,
    confirmed_by: str | None = None,
):
    if not prepared_by and not confirmed_by:
        return

    if prepared_by:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"Prepared by: {prepared_by}.")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(11)

    if confirmed_by:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"Approved by: {confirmed_by}.")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(11)


def build_final_report(
    wideload_df: pd.DataFrame,
    impounded_prohibited_df: pd.DataFrame,
    overloaded_df: pd.DataFrame,
    report_date: str,
    station: str,
    bound: str,
    daily_df: pd.DataFrame | None = None,
    prepared_by: str | None = None,
    confirmed_by: str | None = None,
    traffic_census: dict | None = None,
    daily_summary: dict | None = None,
    transgressions: dict | None = None,
    wideload_count: int | None = None,
) -> io.BytesIO:
    context = ReportContext(
        report_date=report_date,
        station=station,
        bound=bound,
        wideload_count=wideload_count if wideload_count is not None else len(wideload_df),
        overloaded_valid_permit_count=count_valid_permit_vehicles(overloaded_df),
    )

    doc = Document()

    apply_standard_layout(
        doc,
        report_date=context.report_date,
        station=context.station,
        bound=context.bound,
    )

    add_report_title(doc, context.station, context.bound)

    totals = get_daily_hour_totals_from_dataframe(daily_df)

    if daily_df is not None:
        add_daily_hour_statistics_section(doc, daily_df)
        add_prepared_confirmed_lines(doc, prepared_by, confirmed_by)

        doc.add_page_break()
        add_daily_hour_chart_section(doc, daily_df)
        doc.add_page_break()

    if traffic_census is not None:
        total_traffic = (
            totals["q"]
            + totals["x"]
            + int(traffic_census["total_traffic_census"])
            + context.wideload_count
        )
        add_traffic_census_section(
            doc,
            traffic_census,
            exemption_not_weighed=context.wideload_count,
            total_weighed=totals["x"],
            hswim_cleared=totals["q"],
            total_traffic=total_traffic,
        )

    if daily_summary is not None:
        add_daily_summary_section(doc, daily_summary)

    if transgressions is not None:
        add_transgressions_section(doc, transgressions)

    add_impounded_prohibited_section(doc, impounded_prohibited_df)

    add_wideload_section(doc, wideload_df)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
