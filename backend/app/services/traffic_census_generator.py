from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.services.report_layout import apply_standard_layout
from app.services.traffic_census_processor import traffic_census_rows

import io


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table):
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)

    tbl_layout.set(qn("w:type"), "fixed")


def set_table_grid(table, widths):
    tbl = table._tbl

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    tbl_grid = OxmlElement("w:tblGrid")

    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    tbl.insert(0, tbl_grid)


def style_cell(
    cell,
    font_size=8,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for paragraph in cell.paragraphs:
        paragraph.alignment = align

        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.bold = bold

            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)

            r_fonts.set(qn("w:ascii"), "Arial")
            r_fonts.set(qn("w:hAnsi"), "Arial")


def add_traffic_census_section(
    doc: Document,
    traffic_census_data: dict,
    exemption_not_weighed: int = 0,
    total_weighed: int = 0,
    hswim_cleared: int = 0,
    total_traffic: int | None = None,
):
    normalized = dict(traffic_census_data)

    buses = int(normalized["buses_gte_3500kg"])
    vehicles_3500_7000 = int(normalized["vehicles_3500_to_7000_excluding_buses"])
    vehicles_7000 = int(normalized["vehicles_gte_7000_excluding_buses"])
    total_census = int(normalized["total_traffic_census"])

    if total_traffic is None:
        total_traffic = (
            total_census
            + exemption_not_weighed
            + total_weighed
            + hswim_cleared
        )

        
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run("3. TRAFFIC CENSUS DATA")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    headers = [
        "Buses>= 3500kg",
        "Vehicles>= 3500kg\nbut <7000 excluding\nbuses",
        "Vehicles>=\n7000\nexcluding\nbuses",
        "Total\nTraffic\nCensus\n(K)",
        "Exemption\npermits Not\nweighed (E)",
        "Total Weighed",
        "Total Traffic",
    ]

    values = [
        f"{buses:,}",
        f"{vehicles_3500_7000:,}",
        f"{vehicles_7000:,}",
        f"{total_census:,}",
        f"{exemption_not_weighed:,}",
        f"{total_weighed:,}",
        f"{total_traffic:,}",
    ]

    table = doc.add_table(rows=2, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    widths = [2260, 2392, 1740, 1322, 1487, 1740, 3549]

    set_fixed_table_layout(table)
    set_table_grid(table, widths)

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        set_cell_width(cell, widths[col_idx])
        style_cell(cell, font_size=11, bold=True)

    table.rows[0].height = Pt(36.4)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for col_idx, value in enumerate(values):
        cell = table.rows[1].cells[col_idx]
        cell.text = value
        set_cell_width(cell, widths[col_idx])
        style_cell(cell, font_size=11, bold=False)

    table.rows[1].height = Pt(20.5)
    table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

def generate_traffic_census_report(
    traffic_census_data: dict,
    report_date: str,
    station: str,
    bound: str,
) -> io.BytesIO:
    doc = Document()

    apply_standard_layout(
        doc,
        report_date=report_date,
        station=station,
        bound=bound,
    )

    add_traffic_census_section(doc, traffic_census_data)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
