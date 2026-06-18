from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pathlib import Path
from datetime import datetime


LOGO_PATH = Path(__file__).resolve().parent.parent / "assets" / "bgwhitelogo.png"
LETTER_LANDSCAPE_WIDTH_INCHES = 11.0
LETTER_LANDSCAPE_HEIGHT_INCHES = 8.5
TOP_MARGIN_INCHES = 0.5
BOTTOM_MARGIN_INCHES = 0.5
LEFT_MARGIN_INCHES = 0.5
RIGHT_MARGIN_INCHES = 0.5
HEADER_DISTANCE_INCHES = 0.28
FOOTER_DISTANCE_INCHES = 0.28
TWIPS_PER_INCH = 1440
LETTER_PRINTABLE_WIDTH_TWIPS = int(
    (LETTER_LANDSCAPE_WIDTH_INCHES - LEFT_MARGIN_INCHES - RIGHT_MARGIN_INCHES)
    * TWIPS_PER_INCH
)
LETTER_TABLE_WIDTH_TWIPS = 15408
A4_LANDSCAPE_WIDTH_INCHES = LETTER_LANDSCAPE_WIDTH_INCHES
A4_LANDSCAPE_HEIGHT_INCHES = LETTER_LANDSCAPE_HEIGHT_INCHES
A4_PRINTABLE_WIDTH_TWIPS = LETTER_TABLE_WIDTH_TWIPS
A4_TABLE_WIDTH_TWIPS = LETTER_TABLE_WIDTH_TWIPS

FONT_NAME = "Arial"

REPORT_TITLE_SIZE = 11
SECTION_TITLE_SIZE = 11
SUBHEADING_SIZE = 11
TABLE_HEADER_SIZE = 10
TABLE_BODY_SIZE = 10
FOOTER_SIZE = 11


def apply_standard_layout(doc, report_date, station, bound):
    enable_field_updates(doc)
    section = doc.sections[0]

    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(LETTER_LANDSCAPE_WIDTH_INCHES)
    section.page_height = Inches(LETTER_LANDSCAPE_HEIGHT_INCHES)

    section.top_margin = Inches(TOP_MARGIN_INCHES)
    section.bottom_margin = Inches(BOTTOM_MARGIN_INCHES)
    section.left_margin = Inches(LEFT_MARGIN_INCHES)
    section.right_margin = Inches(RIGHT_MARGIN_INCHES)

    section.header_distance = Inches(HEADER_DISTANCE_INCHES)
    section.footer_distance = Inches(FOOTER_DISTANCE_INCHES)

    add_header(section)
    add_footer(section, report_date, station, bound)


def enable_field_updates(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))

    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)

    update_fields.set(qn("w:val"), "true")


def add_header(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    if LOGO_PATH.exists():
        run = paragraph.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.5))


def get_report_date_parts(date_str):
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        day = dt.day
        suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
        day_str = str(day)
        month_year_str = dt.strftime(" %B %Y")
        return day_str, suffix, month_year_str
    except Exception:
        return date_str, "", ""


def style_footer_run(run):
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(FOOTER_SIZE)


def add_field(paragraph, field_name):
    run = paragraph.add_run()
    style_footer_run(run)

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field_name} "

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")

    field_text = OxmlElement("w:t")
    field_text.text = "1"

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_separate)
    run._r.append(field_text)
    run._r.append(field_end)


def add_footer(section, report_date, station, bound):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    day_str, suffix, month_year_str = get_report_date_parts(report_date)

    station_name = station.upper()
    if "WEIGHBRIDGE" not in station_name:
        station_name = f"{station_name} WEIGHBRIDGE"
    bound_name = bound.upper()
    if "BOUND" not in bound_name:
        bound_name = f"{bound_name} BOUND"

    # Prefix run up to day number
    prefix = (
        "KeNHA/WB/MTCE/4339/2025        "
        f"{station_name} {bound_name} DAILY REPORT {day_str}"
    )
    run_prefix = paragraph.add_run(prefix)
    style_footer_run(run_prefix)

    # Superscripted suffix run
    if suffix:
        run_suffix = paragraph.add_run(suffix)
        style_footer_run(run_suffix)
        run_suffix.font.superscript = True

    # Mid run containing month/year and page prefix
    mid_text = f"{month_year_str}        Page "
    run_mid = paragraph.add_run(mid_text)
    style_footer_run(run_mid)

    add_field(paragraph, "PAGE")
    run_end = paragraph.add_run(" of ")
    style_footer_run(run_end)
    add_field(paragraph, "NUMPAGES")
