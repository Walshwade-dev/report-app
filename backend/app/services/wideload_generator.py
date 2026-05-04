from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
)
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.services.report_layout import apply_standard_layout

import io
import pandas as pd


HEADER_LABELS = {
    "Inspection Date": "Inspection\nDate",
    "Permit issue date": "Permit issue\ndate",
    "Abnormal Load Permit": "Abnormal\nLoad Permit",
    "Authorized Weight": "Authorized\nWeight",
    "Date of Travel": "Date\nof Travel",
}

COLUMN_RATIOS = {
    "Cargo": 3,
    "Transporter": 2,
    "Permit issue date": 1.5,
}


def get_column_widths(columns, total_width=16200):
    total_ratio = sum(COLUMN_RATIOS.get(col, 1) for col in columns)
    base_width = total_width / total_ratio

    return {
        col: int(base_width * COLUMN_RATIOS.get(col, 1))
        for col in columns
    }


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table):
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)

    tbl_layout.set(qn("w:type"), "fixed")


def set_table_grid(table, columns, widths):
    tbl = table._tbl

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    tbl_grid = OxmlElement("w:tblGrid")

    for col in columns:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(widths[col]))
        tbl_grid.append(grid_col)

    tbl.insert(0, tbl_grid)


def apply_widths_to_all_cells(table, columns, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[columns[i]])


def set_cell_text_direction(cell, direction="btLr"):
    tc_pr = cell._tc.get_or_add_tcPr()

    text_direction = tc_pr.find(qn("w:textDirection"))
    if text_direction is None:
        text_direction = OxmlElement("w:textDirection")
        tc_pr.append(text_direction)

    text_direction.set(qn("w:val"), direction)


def style_cell(
    cell,
    font_size=6,
    bold=False,
    vertical=False,
    valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
    align=WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.vertical_alignment = valign

    if vertical:
        set_cell_text_direction(cell)

    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.bold = bold


def add_wideload_section(doc: Document, df: pd.DataFrame):
    heading = doc.add_paragraph()
    run = heading.add_run("7. VEHICLE INSPECTION REPORT (WIDE LOADS)")
    run.bold = True
    run.underline = True
    run.font.size = Pt(10)

    columns = list(df.columns)
    widths = get_column_widths(columns)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    set_fixed_table_layout(table)
    set_table_grid(table, columns, widths)

    header_row = table.rows[0]
    header_row.height = Inches(0.8)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for i, col in enumerate(columns):
        cell = header_row.cells[i]
        cell.text = HEADER_LABELS.get(col, col)

        style_cell(
            cell,
            font_size=6,
            bold=True,
            vertical=True,
            valign=WD_CELL_VERTICAL_ALIGNMENT.TOP,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

    for _, row in df.iterrows():
        row_obj = table.add_row()
        row_obj.height = Inches(0.8)
        row_obj.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        for i, value in enumerate(row):
            cell = row_obj.cells[i]
            cell.text = "" if pd.isna(value) else str(value).upper()

            style_cell(
                cell,
                font_size=6,
                vertical=True,
                valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

    apply_widths_to_all_cells(table, columns, widths)

def generate_wideload_report(
    df: pd.DataFrame,
    report_date: str,
    station: str,
    bound: str,
) -> io.BytesIO:
    doc = Document()

    apply_standard_layout(
        doc,
        report_date=report_date,
        station=station,
        bound=bound,
    )

    add_wideload_section(doc, df)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer