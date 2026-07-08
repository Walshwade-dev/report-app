from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.services.report_layout import A4_TABLE_WIDTH_TWIPS
from app.services.transgressions_processor import (
    ACTION_REPORT_COLUMNS,
    DAILY_TRANSGRESSIONS_COLUMNS,
    normalize_transgressions_input,
)


DAILY_WIDTH_RATIOS = {
    "Date": 1.0,
    "Time": 0.9,
    "Reg No": 1.0,
    "Axle Config": 1.0,
    "Transporter": 1.4,
    "Census Clerk": 1.2,
    "Police In charge": 1.4,
    "Action Taken": 1.5,
    "Caught": 0.8,
    "Next WB report sent": 1.3,
    "Next WB": 1.0,
}

ACTION_WIDTH_RATIOS = {
    "Date": 1.0,
    "Time Received": 1.1,
    "Truck No.": 1.0,
    "Sending WB station": 1.4,
    "OCS Reported To": 1.4,
    "Action 1": 1.2,
    "Action 2": 1.2,
    "Attach evidence if any": 1.5,
    "Weight Noted": 1.1,
    "Tagged in System": 1.2,
}


def get_widths(columns, ratios, total_width=A4_TABLE_WIDTH_TWIPS):
    if columns == DAILY_TRANSGRESSIONS_COLUMNS:
        return [1368, 1169, 1442, 988, 1712, 1620, 1733, 1415, 991, 1350, 1047]

    if columns == ACTION_REPORT_COLUMNS:
        return [1349, 1197, 1323, 1251, 1343, 1366, 1371, 2831, 1556, 1173]

    ratio_total = sum(ratios.get(column, 1) for column in columns)
    base_width = total_width / ratio_total
    return [int(base_width * ratios.get(column, 1)) for column in columns]


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table):
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)

    tbl_layout.set(qn("w:type"), "fixed")


def set_table_grid(table, widths):
    tbl = table._tbl

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    tbl_grid = OxmlElement("w:tblGrid")

    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    tbl.insert(0, tbl_grid)


def style_cell(
    cell,
    font_size=10,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for paragraph in cell.paragraphs:
        paragraph.alignment = align

        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.bold = bold

            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)

            r_fonts.set(qn("w:ascii"), "Arial")
            r_fonts.set(qn("w:hAnsi"), "Arial")


def add_subheading(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_table(doc: Document, columns, rows, ratios):
    widths = get_widths(columns, ratios)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    set_fixed_table_layout(table)
    set_table_grid(table, widths)

    header_row = table.rows[0]
    header_row.height = Pt(31.2)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for index, column in enumerate(columns):
        cell = header_row.cells[index]
        cell.text = column
        set_cell_width(cell, widths[index])
        style_cell(cell, font_size=11, bold=True)

    if not rows:
        row = table.add_row()
        row.height = Pt(41.8)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        for index, column in enumerate(columns):
            cell = row.cells[index]

            if index == 0:
                cell.text = "NIL"
            elif column == "Next WB":
                cell.text = "-"
            else:
                cell.text = ""

            set_cell_width(cell, widths[index])
            style_cell(cell, font_size=10, bold=False)

        return table

    for source_row in rows:
        row = table.add_row()
        row.height = Pt(20)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        for index, column in enumerate(columns):
            cell = row.cells[index]
            cell.text = source_row.get(column, "")
            set_cell_width(cell, widths[index])
            style_cell(cell, font_size=10)

    return table


def add_transgressions_section(doc: Document, transgressions_data: dict):
    normalized = normalize_transgressions_input(transgressions_data)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run("5. TRANSGRESSIONS")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    add_subheading(doc, "DAILY TRANSGRESSIONS REPORT")
    add_table(
        doc,
        DAILY_TRANSGRESSIONS_COLUMNS,
        normalized["daily_transgressions"],
        DAILY_WIDTH_RATIOS,
    )

    add_subheading(doc, "TRANSGRESSIONS ACTION REPORT")
    add_table(
        doc,
        ACTION_REPORT_COLUMNS,
        normalized["action_report"],
        ACTION_WIDTH_RATIOS,
    )
